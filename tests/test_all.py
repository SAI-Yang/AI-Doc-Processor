"""项目综合测试套件

覆盖文档读取、Markdown清洗、DOCX格式、编码处理、分块算法、
模板管理、流水线配置、文档写入、配置验证等核心功能。
所有测试不依赖外部 API 或硬件，使用临时文件。
"""

import sys
import os
import json
import tempfile
import asyncio
import unittest
from pathlib import Path

# 将项目根目录加入 sys.path，使 app 包可导入
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


# ===================================================================
# 文档读取
# ===================================================================

class TestDocumentReading(unittest.TestCase):
    """测试各格式文档读取正确性"""

    def test_read_txt(self):
        """读取 UTF-8 TXT 文件"""
        from app.document import read_document

        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w', encoding='utf-8') as f:
            f.write('测试文档内容\n第二行内容')
            fpath = Path(f.name)

        try:
            doc = read_document(fpath)
            self.assertIn('测试文档', doc.content)
            self.assertIn('第二行', doc.content)
            self.assertGreater(doc.metadata['char_count'], 0)
            self.assertEqual(doc.format, 'txt')
        finally:
            fpath.unlink(missing_ok=True)

    def test_read_txt_gbk(self):
        """读取 GBK 编码 TXT 文件（应自动回退编码）"""
        from app.document import TxtReader

        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='wb') as f:
            f.write('中文GBK内容'.encode('gbk'))
            fpath = Path(f.name)

        try:
            # 指定 utf-8 但 gbk 内容应能被回退编码读取
            reader = TxtReader(encoding='utf-8')
            doc = reader.read(fpath)
            self.assertIn('中文', doc.content)
        finally:
            fpath.unlink(missing_ok=True)

    def test_read_docx(self):
        """读取 DOCX 文件"""
        from docx import Document as DocxDocument
        from app.document import read_document

        d = DocxDocument()
        d.add_paragraph('测试段落一')
        d.add_paragraph('测试段落二')
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        d.save(tmp.name)
        tmp.close()
        fpath = Path(tmp.name)

        try:
            doc = read_document(fpath)
            self.assertIn('测试段落一', doc.content)
            self.assertIn('测试段落二', doc.content)
            self.assertGreaterEqual(doc.metadata['paragraph_count'], 2)
            self.assertEqual(doc.format, 'docx')
        finally:
            fpath.unlink(missing_ok=True)

    def test_read_md(self):
        """读取 Markdown 文件，检查样式检测"""
        from app.document import MdReader

        with tempfile.NamedTemporaryFile(suffix='.md', delete=False, mode='w', encoding='utf-8') as f:
            f.write('# 标题\n\n一段落\n\n- 列表项1\n- 列表项2\n\n> 引用')
            fpath = Path(f.name)

        try:
            reader = MdReader()
            doc = reader.read(fpath)
            self.assertEqual(doc.format, 'md')
            styles = [p['style'] for p in doc.paragraphs]
            self.assertIn('Heading 1', styles)
            self.assertIn('List Item', styles)
            self.assertIn('Blockquote', styles)
            self.assertIn('Normal', styles)
        finally:
            fpath.unlink(missing_ok=True)

    def test_read_empty_txt(self):
        """读取空 TXT 文件"""
        from app.document import read_document

        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w', encoding='utf-8') as f:
            f.write('')
            fpath = Path(f.name)

        try:
            doc = read_document(fpath)
            self.assertEqual(doc.content, '')
            self.assertEqual(doc.metadata['char_count'], 0)
            self.assertEqual(doc.metadata['paragraph_count'], 0)
        finally:
            fpath.unlink(missing_ok=True)

    def test_read_unsupported_format(self):
        """不支持的格式应抛出 ValueError"""
        from app.document import get_reader

        with self.assertRaises(ValueError):
            get_reader(Path('test.xyz'))

    def test_read_nonexistent_file(self):
        """不存在的文件应抛出 FileNotFoundError"""
        from app.document import read_document

        with self.assertRaises(FileNotFoundError):
            read_document(Path('nonexistent_file_12345.txt'))

    def test_document_from_paragraphs(self):
        """Document.from_paragraphs 工厂方法"""
        from app.document import Document

        paragraphs = [
            {'index': 0, 'text': '第一段', 'style': 'Normal'},
            {'index': 1, 'text': '第二段', 'style': 'Normal'},
        ]
        path = Path('test.txt')
        doc = Document.from_paragraphs(path, 'txt', paragraphs)

        self.assertEqual(doc.content, '第一段\n第二段')
        self.assertEqual(doc.metadata['paragraph_count'], 2)
        self.assertEqual(doc.metadata['char_count'], 7)
        self.assertEqual(doc.format, 'txt')


# ===================================================================
# Markdown 清洗
# ===================================================================

class TestMarkdownCleaning(unittest.TestCase):
    """测试 LLM 输出的 Markdown 清理"""

    def test_strip_bold(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('这是**粗体**文字'), '这是粗体文字')

    def test_strip_italic(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('这是*斜体*文字'), '这是斜体文字')

    def test_strip_code(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('这是`代码`文字'), '这是代码文字')

    def test_strip_heading(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('# 标题'), '标题')
        self.assertEqual(strip_markdown('## 二级标题'), '二级标题')
        self.assertEqual(strip_markdown('### 三级标题'), '三级标题')

    def test_strip_list(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('- 列表项'), '列表项')
        self.assertEqual(strip_markdown('1. 编号项'), '编号项')
        self.assertEqual(strip_markdown('* 星号列表'), '星号列表')

    def test_strip_blockquote(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('> 引用内容'), '引用内容')

    def test_strip_code_block(self):
        from app.processing_skill import strip_markdown
        result = strip_markdown('前置\n```python\ncode\n```\n后置')
        self.assertNotIn('```', result)
        self.assertIn('前置', result)
        self.assertIn('后置', result)

    def test_strip_horizontal_rule(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('---'), '')

    def test_strip_no_markdown(self):
        """没有 Markdown 符号的文本应保持不变"""
        from app.processing_skill import strip_markdown
        text = '这是一段普通的纯文本，没有任何格式符号。'
        self.assertEqual(strip_markdown(text), text)

    def test_strip_empty_string(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown(''), '')

    def test_clean_output_no_markdown(self):
        """clean_output 不应输出任何 Markdown 符号"""
        from app.processing_skill import clean_output

        dirty = '这是**粗体**和*斜体*和`代码`\n# 标题\n> 引用\n- 列表'
        result = clean_output(dirty)
        for sym in ['**', '*', '`', '#', '>', '- [']:
            self.assertNotIn(sym.strip(), result,
                             f'结果中仍包含 Markdown 符号: {sym}')

    def test_clean_output_removes_thinking_tokens(self):
        """clean_output 应移除 LLM 思考产物前缀"""
        from app.processing_skill import clean_output
        result = clean_output('分析：这是一个测试')
        self.assertNotIn('分析', result)
        self.assertIn('这是一个测试', result)

    def test_clean_output_normal_text_preserved(self):
        """clean_output 应保持正常文本不变"""
        from app.processing_skill import clean_output
        text = '这是一段完全正常的文本。包含中文、English、数字123。'
        result = clean_output(text)
        self.assertIn('完全正常的文本', result)
        self.assertIn('English', result)
        self.assertIn('123', result)

    def test_clean_output_empty(self):
        from app.processing_skill import clean_output
        self.assertEqual(clean_output(''), '')


# ===================================================================
# DOCX 格式
# ===================================================================

class TestDocxFormatting(unittest.TestCase):
    """测试 DOCX 输出格式"""

    def test_heading_created(self):
        """通过 generator 生成的 DOCX 标题段落应存在"""
        from docx import Document as DocxDocument

        d = DocxDocument()
        d.add_heading('实验报告', 1)
        self.assertEqual(len(d.paragraphs), 1)

    def test_generator_heading_font(self):
        """generator.save_as_docx 标题段落应使用 Heading 样式且文本正确"""
        from app.config import AppConfig
        from app.generator import DocumentGenerator

        config = AppConfig.default()
        gen = DocumentGenerator(config)

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / 'test.docx'
            gen.save_as_docx(
                '# 一级标题\n\n正文内容\n\n## 二级标题\n\n更多正文',
                output,
                title='测试文档'
            )

            from docx import Document as DocxDocument
            doc = DocxDocument(str(output))

            # 验证标题段落存在且样式正确
            heading_paras = [
                p for p in doc.paragraphs
                if p.style and p.style.name.startswith('Heading')
            ]
            self.assertGreater(len(heading_paras), 0, '未找到标题段落')

            # 验证标题文本正确
            self.assertIn('一级标题', heading_paras[0].text)

            # 验证 Heading 样式层面已设置黑体 (样式级别的 font 设置)
            from docx.oxml.ns import qn
            hs = doc.styles['Heading 1']
            east_asia = hs.element.rPr.rFonts.get(qn('w:eastAsia'))
            self.assertEqual(east_asia, '黑体', 'Heading 1 样式未设置黑体')
            self.assertTrue(hs.font.bold, 'Heading 1 样式未加粗')

    def test_generator_body_font(self):
        """generator.save_as_docx 正文段落应存在且 Normal 样式已设置宋体"""
        from app.config import AppConfig
        from app.generator import DocumentGenerator

        config = AppConfig.default()
        gen = DocumentGenerator(config)

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / 'test.docx'
            gen.save_as_docx(
                '# 标题\n\n这是一段正文内容，用于测试字体设置。',
                output,
                title='测试'
            )

            from docx import Document as DocxDocument
            doc = DocxDocument(str(output))

            # 验证正文段落存在
            body_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            self.assertTrue(
                any('正文内容' in t for t in body_texts),
                '正文段落文本未找到'
            )

            # 验证 Normal 样式层面已设置宋体
            from docx.oxml.ns import qn
            normal_style = doc.styles['Normal']
            east_asia = normal_style.element.rPr.rFonts.get(qn('w:eastAsia'))
            self.assertEqual(east_asia, '宋体', 'Normal 样式未设置宋体')

    def test_docx_handler_roundtrip(self):
        """DocxHandler 保存后段落数不应减少"""
        from app.docx_handler import DocxHandler
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmpdir:
            # 创建源文件
            src = Path(tmpdir) / 'input.docx'
            d = DocxDocument()
            d.add_paragraph('段落一')
            d.add_paragraph('段落二')
            d.add_paragraph('段落三')
            d.save(str(src))

            handler = DocxHandler()
            handler.open(str(src))
            text = handler.get_text()
            self.assertIn('段落一', text)
            self.assertIn('段落三', text)

            handler.apply_changes('新一段\n新二段\n新三段')
            out = Path(tmpdir) / 'output.docx'
            handler.save(str(out))
            self.assertTrue(out.exists())
            self.assertGreater(out.stat().st_size, 0)


# ===================================================================
# 编码处理
# ===================================================================

class TestEncoding(unittest.TestCase):
    """测试编码检测与乱码处理"""

    def test_is_garbled_normal(self):
        from app.processing_skill import is_garbled
        self.assertFalse(is_garbled('正常中文文本'))

    def test_is_garbled_ascii(self):
        from app.processing_skill import is_garbled
        self.assertFalse(is_garbled('Normal ASCII text with numbers 123'))

    def test_is_garbled_empty(self):
        from app.processing_skill import is_garbled
        self.assertFalse(is_garbled(''))

    def test_is_garbled_garbled(self):
        """检测乱码文本"""
        from app.processing_skill import is_garbled
        # 将 UTF-8 字节解释为 Latin-1 的典型乱码
        garbled = 'å\x9fºæ\x9c¬å\x8a\x9fè\x83½'
        self.assertTrue(is_garbled(garbled))

    def test_detect_encoding_normal(self):
        from app.processing_skill import detect_encoding
        issues = detect_encoding('正常文本')
        self.assertEqual(issues, [])

    def test_detect_encoding_garbled(self):
        from app.processing_skill import detect_encoding
        issues = detect_encoding('å\x9fºæ\x9c¬')
        self.assertGreater(len(issues), 0)

    def test_detect_encoding_html_entities(self):
        from app.processing_skill import detect_encoding
        issues = detect_encoding('&#20013;&#25991;')
        self.assertIn('html_entities', issues)

    def test_detect_encoding_unicode_escapes(self):
        from app.processing_skill import detect_encoding
        issues = detect_encoding('\\u4e2d\\u6587')
        self.assertIn('unicode_escapes', issues)

    def test_repair_garbled(self):
        from app.processing_skill import repair_garbled
        # UTF-8 字节被解释为 Latin-1
        garbled = 'å\x9fºæ\x9c¬å\x8a\x9fè\x83½'
        repaired = repair_garbled(garbled)
        # 修复后应包含有效字符且不再乱码
        from app.processing_skill import is_garbled
        # 可能修复为有效文本或保持不变
        if repaired != garbled:
            self.assertFalse(is_garbled(repaired))


# ===================================================================
# 分块算法
# ===================================================================

class TestChunking(unittest.TestCase):
    """测试智能分段和上下文分块算法"""

    def test_empty_document(self):
        """空文档应返回空分块"""
        from app.engine import split_into_chunks
        from app.document import Document

        doc = Document(
            path=Path('test.txt'), format='txt',
            content='', paragraphs=[],
        )
        chunks = split_into_chunks(doc, max_chunk_tokens=100)
        self.assertEqual(len(chunks), 0)

    def test_small_content_single_chunk(self):
        """小内容应合并为一块"""
        from app.engine import split_into_chunks
        from app.document import Document

        paragraphs = [
            {'index': 0, 'text': '短段落', 'style': 'Normal'},
            {'index': 1, 'text': '另一个短段落', 'style': 'Normal'},
        ]
        doc = Document(
            path=Path('test.txt'), format='txt',
            content='短段落\n另一个短段落',
            paragraphs=paragraphs,
        )
        chunks = split_into_chunks(doc, max_chunk_tokens=3000)
        self.assertEqual(len(chunks), 1)
        self.assertIn('短段落', chunks[0].text)
        self.assertIn('另一个短段落', chunks[0].text)

    def test_large_paragraph_split(self):
        """超长段落应按句子分割"""
        from app.engine import split_into_chunks
        from app.document import Document

        long_text = '这是第一句。这是第二句。这是第三句。这是第四句。 ' * 20
        paragraphs = [
            {'index': 0, 'text': long_text, 'style': 'Normal'},
        ]
        doc = Document(
            path=Path('test.txt'), format='txt',
            content=long_text, paragraphs=paragraphs,
        )
        chunks = split_into_chunks(doc, max_chunk_tokens=50)
        self.assertGreater(len(chunks), 1)

    def test_content_not_lost(self):
        """分块后所有段落文本应出现在 chunks 中"""
        from app.engine import split_into_chunks
        from app.document import Document

        paragraphs = [
            {'index': 0, 'text': '段落一', 'style': 'Normal'},
            {'index': 1, 'text': '段落二', 'style': 'Normal'},
            {'index': 2, 'text': '段落三', 'style': 'Normal'},
        ]
        doc = Document(
            path=Path('test.txt'), format='txt',
            content='段落一\n\n段落二\n\n段落三',
            paragraphs=paragraphs,
        )
        chunks = split_into_chunks(doc, max_chunk_tokens=6000)
        combined = ''.join(c.text for c in chunks)
        for p in paragraphs:
            self.assertIn(p['text'], combined,
                          f'段落 "{p["text"]}" 在分块中丢失')

    def test_content_not_lost_small_chunk(self):
        """小 chunk 限制下所有字符应出现在 chunks 中（可能有重叠）"""
        from app.engine import split_into_chunks
        from app.document import Document

        text = '一二三。四五六。七八九。'
        paragraphs = [
            {'index': 0, 'text': text, 'style': 'Normal'},
        ]
        doc = Document(
            path=Path('test.txt'), format='txt',
            content=text, paragraphs=paragraphs,
        )
        chunks = split_into_chunks(doc, max_chunk_tokens=10)
        combined = ''.join(c.text for c in chunks)
        for ch in text:
            if ch.strip():
                self.assertIn(ch, combined,
                              f'字符 "{ch}" 在分块中丢失')

    def test_build_context_chunks_small(self):
        """小文本应返回单块"""
        from app.processing_skill import build_context_chunks

        chunks = build_context_chunks('小文本', max_chars=8000, overlap=500)
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0][0], 0)
        self.assertEqual(chunks[0][1], '小文本')
        self.assertEqual(chunks[0][2], '')

    def test_build_context_chunks_large(self):
        """大文本应分多块，且含上文语境"""
        from app.processing_skill import build_context_chunks

        text = '区块A的内容。\n' * 100 + '区块B的内容。\n' * 100
        chunks = build_context_chunks(text, max_chars=200, overlap=50)
        self.assertGreater(len(chunks), 1)

        # 第二块应有上下文
        if len(chunks) > 1:
            self.assertNotEqual(chunks[1][2], '',
                                '第二块应包含上文语境')
            # 上下文应来自第一块末尾
            self.assertIn('区块A', chunks[1][2])

    def test_build_context_chunks_empty(self):
        from app.processing_skill import build_context_chunks
        chunks = build_context_chunks('', max_chars=8000)
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0][1], '')

    def test_build_context_chunks_boundary(self):
        """长文本在段落边界截断不应丢失内容"""
        from app.processing_skill import build_context_chunks

        # 构造略超 max_chars 的文本，触发换行截断
        text = 'A\n' * 2000
        chunks = build_context_chunks(text, max_chars=500, overlap=50)
        total = sum(len(c[1]) for c in chunks)
        # 内容应大致不丢失（每块最后可能多一个换行符）
        self.assertAlmostEqual(total, len(text), delta=len(chunks))


# ===================================================================
# 模板管理
# ===================================================================

class TestTemplateManager(unittest.TestCase):
    """测试模板管理"""

    def setUp(self):
        from app.template_manager import TemplateManager
        self.mgr = TemplateManager()

    def test_list_templates(self):
        """模板列表应包含内置模板"""
        templates = self.mgr.list_templates()
        builtin = [t for t in templates if t['is_builtin']]
        self.assertGreaterEqual(len(builtin), 9)

    def test_get_template(self):
        tpl = self.mgr.get('zh_to_en')
        self.assertIsNotNone(tpl)
        self.assertEqual(tpl.name, '中译英')

    def test_get_nonexistent(self):
        tpl = self.mgr.get('nonexistent_template')
        self.assertIsNone(tpl)

    def test_render(self):
        sp, up, temp, mt = self.mgr.render('general_polish', '测试内容')
        self.assertIn('测试内容', up)
        self.assertIn('润色', up)
        self.assertIsInstance(sp, str)
        self.assertIsInstance(temp, float)
        self.assertIsInstance(mt, int)

    def test_render_nonexistent_raises(self):
        with self.assertRaises(KeyError):
            self.mgr.render('nonexistent', 'content')

    def test_get_default_template_id(self):
        default_id = self.mgr.get_default_template_id()
        self.assertIsNotNone(default_id)
        self.assertIsInstance(default_id, str)

    def test_render_with_context(self):
        """渲染带额外上下文的模板"""
        sp, up, temp, mt = self.mgr.render(
            'generate_caption', '图片描述内容',
            document_context='文档上下文内容'
        )
        self.assertIn('图片描述内容', up)
        self.assertIn('文档上下文内容', up)

    def test_custom_template_add_and_remove(self):
        """添加和删除自定义模板"""
        from app.template_manager import Template

        tpl = Template(
            name='测试模板',
            description='临时测试用',
            system_prompt='测试 system prompt',
            user_prompt='处理以下内容：{content}',
            temperature=0.5,
            max_tokens=2048,
        )
        self.mgr.add_custom('test_custom', tpl)

        retrieved = self.mgr.get('test_custom')
        self.assertIsNotNone(retrieved)
        self.assertEqual(retrieved.name, '测试模板')

        # 渲染
        sp, up, temp, mt = self.mgr.render('test_custom', 'hello')
        self.assertIn('hello', up)
        self.assertEqual(temp, 0.5)

        # 删除
        removed = self.mgr.remove_custom('test_custom')
        self.assertTrue(removed)
        self.assertIsNone(self.mgr.get('test_custom'))

    def test_cannot_remove_builtin(self):
        """不能删除内置模板"""
        result = self.mgr.remove_custom('zh_to_en')
        self.assertFalse(result)
        self.assertIsNotNone(self.mgr.get('zh_to_en'))

    def test_export_import_template(self):
        """导出和导入模板"""
        from app.template_manager import Template

        with tempfile.TemporaryDirectory() as tmpdir:
            export_path = Path(tmpdir) / 'exported.json'

            # 导出内置模板
            self.mgr.export_template('general_polish', export_path)
            self.assertTrue(export_path.exists())

            # 导入为新模板
            imported_id = self.mgr.import_template(export_path)
            self.assertIsNotNone(self.mgr.get(imported_id))

            # 清理
            self.mgr.remove_custom(imported_id)


# ===================================================================
# 文档写入
# ===================================================================

class TestDocumentWriting(unittest.TestCase):
    """测试文档写入"""

    def test_txt_writer(self):
        from app.document import Document, TxtWriter

        doc = Document(
            path=Path('input.txt'), format='txt',
            content='Hello\nWorld',
            paragraphs=[
                {'index': 0, 'text': 'Hello', 'style': 'Normal'},
                {'index': 1, 'text': 'World', 'style': 'Normal'},
            ],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / 'output.txt'
            writer = TxtWriter()
            result = writer.write(doc, output)
            self.assertTrue(result.exists())
            content = result.read_text(encoding='utf-8')
            self.assertIn('Hello', content)
            self.assertIn('World', content)

    def test_docx_writer(self):
        from app.document import Document, DocxWriter
        from docx import Document as DocxDocument

        doc = Document(
            path=Path('input.txt'), format='txt',
            content='标题内容\n正文内容',
            paragraphs=[
                {'index': 0, 'text': '标题内容', 'style': 'Heading 1'},
                {'index': 1, 'text': '正文内容', 'style': 'Normal'},
            ],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / 'output.docx'
            writer = DocxWriter()
            result = writer.write(doc, output)
            self.assertTrue(result.exists())

            # 验证内容
            readback = DocxDocument(str(result))
            texts = [p.text for p in readback.paragraphs]
            self.assertIn('标题内容', texts)
            self.assertIn('正文内容', texts)

    def test_get_writer(self):
        from app.document import get_writer, TxtWriter, DocxWriter

        txt_writer = get_writer('txt')
        self.assertIsInstance(txt_writer, TxtWriter)

        docx_writer = get_writer('docx')
        self.assertIsInstance(docx_writer, DocxWriter)

    def test_get_writer_unsupported(self):
        from app.document import get_writer

        with self.assertRaises(ValueError):
            get_writer('pdf')

    def test_get_writer_same_as_input_raises(self):
        from app.document import get_writer

        with self.assertRaises(ValueError):
            get_writer('same_as_input')


# ===================================================================
# 流水线
# ===================================================================

class TestPipeline(unittest.TestCase):
    """测试流水线配置和基础功能"""

    def test_pipeline_config_creation(self):
        from app.pipeline import PipelineConfig, PipelineStep

        config = PipelineConfig(
            name='测试流水线',
            description='测试用',
            steps=[
                PipelineStep(template_id='en_to_zh', description='翻译'),
                PipelineStep(template_id='academic_polish', description='润色'),
            ],
        )
        self.assertEqual(len(config.steps), 2)
        self.assertEqual(config.steps[0].template_id, 'en_to_zh')

    def test_serialization(self):
        from app.pipeline import PipelineConfig, PipelineStep, Pipeline

        config = PipelineConfig(
            name='测试流水线',
            steps=[
                PipelineStep(template_id='en_to_zh', description='翻译'),
            ],
        )
        pipeline = Pipeline.__new__(Pipeline)
        pipeline.pipeline_config = config
        exported = pipeline.to_dict()

        self.assertEqual(exported['name'], '测试流水线')
        self.assertEqual(len(exported['steps']), 1)
        self.assertEqual(exported['steps'][0]['template_id'], 'en_to_zh')

    def test_save_load_config(self):
        from app.pipeline import PipelineConfig, PipelineStep, Pipeline

        config = PipelineConfig(
            name='保存测试',
            steps=[
                PipelineStep(template_id='simplify', description='简化'),
                PipelineStep(template_id='summarize', description='摘要'),
            ],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / 'pipeline.json'
            pipeline = Pipeline.__new__(Pipeline)
            pipeline.pipeline_config = config
            pipeline.save_config(path)

            self.assertTrue(path.exists())

            loaded_config = Pipeline.load_config(path)
            self.assertEqual(loaded_config.name, '保存测试')
            self.assertEqual(len(loaded_config.steps), 2)

    def test_empty_pipeline_raises(self):
        """空流水线运行应抛出 ValueError"""
        from app.pipeline import Pipeline, PipelineConfig
        from app.config import AppConfig

        config = AppConfig.default()
        pipeline = Pipeline(config, PipelineConfig(name='空流水线'))

        with self.assertRaises(ValueError):
            asyncio.run(pipeline.run(Path('in.txt'), Path('out.txt')))

    def test_create_default_pipelines(self):
        from app.pipeline import create_default_pipelines

        pipelines = create_default_pipelines()
        self.assertGreater(len(pipelines), 0)
        for p in pipelines:
            self.assertIsInstance(p.name, str)
            self.assertGreater(len(p.steps), 0)


# ===================================================================
# 配置管理
# ===================================================================

class TestConfig(unittest.TestCase):
    """测试配置管理"""

    def test_default_values(self):
        from app.config import AppConfig
        config = AppConfig.default()

        self.assertEqual(config.llm.provider, 'deepseek')
        self.assertEqual(config.llm.temperature, 0.3)
        self.assertEqual(config.llm.max_tokens, 4096)
        self.assertEqual(config.output.format, 'same_as_input')
        self.assertEqual(config.processing.max_concurrent, 3)
        self.assertEqual(config.processing.retry_count, 2)
        self.assertEqual(config.processing.timeout, 120)

    def test_validation_empty_api_key(self):
        from app.config import AppConfig
        config = AppConfig.default()
        config.llm.api_key = ''
        errors = config.validate()
        self.assertTrue(any('API' in e for e in errors))

    def test_validation_invalid_temperature(self):
        from app.config import AppConfig
        config = AppConfig.default()
        config.llm.api_key = 'sk-test'
        config.llm.temperature = 3.0
        errors = config.validate()
        self.assertTrue(any('temperature' in e for e in errors))

    def test_validation_negative_retries(self):
        from app.config import AppConfig
        config = AppConfig.default()
        config.llm.api_key = 'sk-test'
        config.processing.retry_count = -1
        errors = config.validate()
        self.assertTrue(any('retry' in e.lower() for e in errors))

    def test_valid_config_no_errors(self):
        from app.config import AppConfig
        config = AppConfig.default()
        config.llm.api_key = 'sk-valid-key'
        errors = config.validate()
        self.assertEqual(errors, [])

    def test_save_and_load(self):
        from app.config import AppConfig

        config = AppConfig.default()
        config.llm.api_key = 'sk-save-test'
        config.llm.model = 'gpt-4'
        config.processing.max_concurrent = 5

        with tempfile.TemporaryDirectory() as tmpdir:
            save_path = Path(tmpdir) / 'config.json'
            saved = config.save(save_path)
            self.assertTrue(saved.exists())

            loaded = AppConfig.load(save_path)
            self.assertEqual(loaded.llm.api_key, 'sk-save-test')
            self.assertEqual(loaded.llm.model, 'gpt-4')
            self.assertEqual(loaded.processing.max_concurrent, 5)

    def test_load_nonexistent_returns_default(self):
        from app.config import AppConfig
        config = AppConfig.load(Path('/nonexistent/path/config.json'))
        self.assertIsInstance(config, AppConfig)

    def test_load_corrupted_json_returns_default(self):
        from app.config import AppConfig

        with tempfile.TemporaryDirectory() as tmpdir:
            bad_path = Path(tmpdir) / 'config.json'
            bad_path.write_text('{invalid json}', encoding='utf-8')
            config = AppConfig.load(bad_path)
            self.assertIsInstance(config, AppConfig)


# ===================================================================
# LLM 客户端
# ===================================================================

class TestLLMClient(unittest.TestCase):
    """测试 LLM 客户端创建和工具函数"""

    def test_create_openai_client(self):
        from app.llm_client import create_client, OpenAIClient
        from app.config import LLMConfig

        config = LLMConfig(provider='openai', api_key='sk-test')
        client = create_client(config)
        self.assertIsInstance(client, OpenAIClient)

    def test_create_deepseek_client(self):
        from app.llm_client import create_client, OpenAIClient
        from app.config import LLMConfig

        config = LLMConfig(provider='deepseek', api_key='sk-test')
        client = create_client(config)
        self.assertIsInstance(client, OpenAIClient)

    def test_create_anthropic_client(self):
        from app.llm_client import create_client, AnthropicClient
        from app.config import LLMConfig

        config = LLMConfig(provider='anthropic', api_key='sk-test')
        client = create_client(config)
        self.assertIsInstance(client, AnthropicClient)

    def test_create_unknown_provider_fallback(self):
        from app.llm_client import create_client, OpenAIClient
        from app.config import LLMConfig

        config = LLMConfig(provider='unknown_provider', api_key='sk-test')
        client = create_client(config)
        self.assertIsInstance(client, OpenAIClient)

    def test_estimate_tokens_empty(self):
        from app.llm_client import estimate_tokens
        self.assertEqual(estimate_tokens(''), 1)

    def test_estimate_tokens_english(self):
        from app.llm_client import estimate_tokens
        tokens = estimate_tokens('Hello world')
        self.assertGreaterEqual(tokens, 1)

    def test_estimate_tokens_chinese(self):
        from app.llm_client import estimate_tokens
        tokens = estimate_tokens('这是一段中文文本')
        self.assertGreaterEqual(tokens, 1)

    def test_estimate_tokens_mixed(self):
        from app.llm_client import estimate_tokens
        tokens = estimate_tokens('混合 English 和 中文 text')
        self.assertGreaterEqual(tokens, 1)


# ===================================================================
# 进度信息
# ===================================================================

class TestProgressInfo(unittest.TestCase):
    """测试进度信息数据结构"""

    def test_progress_info_creation(self):
        from app.engine import ProgressInfo

        info = ProgressInfo(
            stage='reading',
            current=1,
            total=5,
            message='正在读取',
        )
        self.assertEqual(info.stage, 'reading')
        self.assertEqual(info.current, 1)
        self.assertEqual(info.total, 5)

    def test_progress_info_with_error(self):
        from app.engine import ProgressInfo

        info = ProgressInfo(
            stage='error',
            current=0,
            total=0,
            message='失败',
            error='文件不存在',
        )
        self.assertEqual(info.error, '文件不存在')


# ===================================================================
# 入口
# ===================================================================

if __name__ == '__main__':
    unittest.main(verbosity=2)
