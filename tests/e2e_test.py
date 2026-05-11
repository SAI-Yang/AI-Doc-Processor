"""
End-to-end functional tests - exercises all major features
No external API dependence, no hardware needed
"""
import sys, os, tempfile, unittest
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from PyQt5.QtWidgets import QApplication
    _qt_app = QApplication.instance() or QApplication(sys.argv)
except:
    _qt_app = None

class TestDocReading(unittest.TestCase):
    def test_txt(self):
        from app.document import read_document
        with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', encoding='utf-8', delete=False) as f:
            f.write('line1\nline2\nline3')
            p = f.name
        doc = read_document(Path(p))
        self.assertIn('line1', doc.content)
        self.assertEqual(doc.metadata['paragraph_count'], 3)
        os.unlink(p)
    def test_gbk(self):
        from app.document import read_document
        with tempfile.NamedTemporaryFile(suffix='.txt', mode='wb', delete=False) as f:
            f.write('test'.encode('gbk'))
            p = f.name
        doc = read_document(Path(p))
        self.assertIn('test', doc.content)
        os.unlink(p)
    def test_docx(self):
        from app.document import read_document
        from docx import Document
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        d = Document()
        d.add_paragraph('docx paragraph')
        d.save(tmp.name); tmp.close()
        doc = read_document(Path(tmp.name))
        self.assertIn('docx paragraph', doc.content)
        os.unlink(tmp.name)

class TestMarkdownStripping(unittest.TestCase):
    def test_bold(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('**bold**'), 'bold')
    def test_italic(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('*italic*'), 'italic')
    def test_code(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('`code`'), 'code')
    def test_heading(self):
        from app.processing_skill import strip_markdown
        self.assertEqual(strip_markdown('# heading'), 'heading')
    def test_list(self):
        from app.processing_skill import strip_markdown
        self.assertIn('item', strip_markdown('- item'))
        self.assertIn('item', strip_markdown('1. item'))
    def test_comprehensive(self):
        from app.processing_skill import clean_output
        dirty = '**bold** and *italic* and `code`\n# heading\n> quote\n- list\nkeep'
        clean = clean_output(dirty)
        for s in ['**', '`', '# ', '> ']:
            self.assertNotIn(s.strip(), clean)
        self.assertIn('keep', clean)

class TestTemplateSystem(unittest.TestCase):
    def test_list(self):
        from app.template_manager import TemplateManager
        self.assertGreaterEqual(len(TemplateManager().list_templates()), 9)
    def test_render(self):
        from app.template_manager import TemplateManager
        sp, up, t, m = TemplateManager().render('general_polish', 'test')
        self.assertIn('test', up)
    def test_no_markdown(self):
        from app.template_manager import TemplateManager
        tm = TemplateManager()
        for t in tm.list_templates():
            tpl = tm.get(t['id'])
            if tpl:
                c = tpl.system_prompt + tpl.user_prompt
                self.assertNotIn('**', c, f'{t["id"]} has **')

class TestDocxFormatting(unittest.TestCase):
    def test_handler_roundtrip(self):
        from docx import Document
        from app.docx_handler import DocxHandler
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        d = Document()
        d.add_heading('Test Heading', 1)
        d.add_paragraph('Test body text.')
        d.save(tmp.name); tmp.close()
        h = DocxHandler()
        h.open(tmp.name)
        h.apply_changes('Modified Heading\nModified body')
        h.save(tmp.name)
        h2 = DocxHandler()
        h2.open(tmp.name)
        self.assertIn('Modified', h2.get_text())
        os.unlink(tmp.name)

class TestEncoding(unittest.TestCase):
    def test_garbled(self):
        from app.processing_skill import is_garbled
        self.assertFalse(is_garbled('normal text'))
        self.assertFalse(is_garbled(''))
        self.assertTrue(is_garbled('\xef\xbf\xbd\xef\xbf\xbd'))

class TestConfig(unittest.TestCase):
    def test_save_load(self):
        from app.config import AppConfig
        cfg = AppConfig.load()
        cfg.font_family = 'Test'
        cfg.save()
        self.assertEqual(AppConfig.load().font_family, 'Test')
        cfg.font_family = 'LXGW WenKai'
        cfg.save()

class TestChunking(unittest.TestCase):
    def test_content_preserved(self):
        from app.engine import split_into_chunks
        from app.document import Document
        doc = Document(Path('t.txt'), 'txt', 'a\nb\nc',
                       paragraphs=[{'index':i,'text':t,'style':'Normal'}
                                  for i,t in enumerate(['a','b','c'])])
        chunks = split_into_chunks(doc, 6000)
        self.assertGreaterEqual(sum(len(c.text) for c in chunks), 5)

class TestImagePlacer(unittest.TestCase):
    def test_parse(self):
        from app.image_placer import ImagePlacer
        r = ImagePlacer().parse_user_instruction('after paragraph 3')
        self.assertIn('type', r)

class TestRefAnalyzer(unittest.TestCase):
    def test_analyze(self):
        from docx import Document
        from app.reference_analyzer import ReferenceAnalyzer
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        d = Document()
        d.add_heading('Title', 1)
        d.add_paragraph('content.')
        d.save(tmp.name); tmp.close()
        r = ReferenceAnalyzer().analyze(tmp.name)
        self.assertIn('title', r)
        os.unlink(tmp.name)

class TestPromptOptimizer(unittest.TestCase):
    def test_optimize(self):
        from app.prompt_optimizer import optimize_prompt
        r = optimize_prompt('translate to english')
        self.assertGreater(len(r), len('translate to english'))

class TestFontManager(unittest.TestCase):
    @unittest.skipIf(_qt_app is None, 'requires Qt')
    def test_load(self):
        from app.font_manager import load_bundled_fonts
        fonts = load_bundled_fonts()
        self.assertGreater(len(fonts), 0)

if __name__ == '__main__':
    unittest.main(verbosity=2)
