"""文档读写抽象层

提供统一的 Document 结构，支持多种格式的读写。
每个 Reader/Writer 负责单一格式，通过注册表统一调度。
"""

from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """统一的文档数据结构"""
    path: Path
    format: str  # docx / pdf / txt / md
    content: str  # 纯文本全文
    metadata: dict = field(default_factory=dict)  # 文件名、页数、字数等
    paragraphs: list[dict] = field(default_factory=list)  # [{index, text, style}]

    @classmethod
    def from_paragraphs(cls, path: Path, fmt: str,
                        paragraphs: list[dict]) -> "Document":
        """从段落列表构建 Document

        Args:
            path: 文件路径
            fmt: 文件格式
            paragraphs: 段落列表，每项含 index, text, style

        Returns:
            Document 实例
        """
        content = "\n".join(p["text"] for p in paragraphs)
        total_chars = len(content)
        word_count = len(content.split())
        metadata = {
            "filename": path.name,
            "size_bytes": path.stat().st_size if path.exists() else 0,
            "char_count": total_chars,
            "word_count": word_count,
            "paragraph_count": len(paragraphs),
        }
        return cls(
            path=path,
            format=fmt,
            content=content,
            metadata=metadata,
            paragraphs=paragraphs,
        )


# ── 读取器基类 ──────────────────────────────────────────────

class BaseReader(ABC):
    """文档读取器抽象基类"""

    @abstractmethod
    def read(self, path: Path) -> Document:
        """读取文档并返回统一的 Document 结构

        Args:
            path: 文件路径

        Returns:
            Document 实例
        """
        ...


# ── .docx 读取器 ───────────────────────────────────────────

class DocxReader(BaseReader):
    """读取 .docx 文件"""

    def read(self, path: Path) -> Document:
        try:
            from docx import Document as DocxDoc
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDoc(str(path))
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal",
                })

        return Document.from_paragraphs(path, "docx", paragraphs)


# ── .pdf 读取器 ───────────────────────────────────────────

class PdfReader(BaseReader):
    """读取 .pdf 文件

    优先使用 pdfplumber（文字提取质量更高），
    若不可用则回退到 PyPDF2。
    """

    def read(self, path: Path) -> Document:
        # 尝试用 pdfplumber
        try:
            return self._read_pdfplumber(path)
        except ImportError:
            pass

        # 回退到 PyPDF2
        try:
            return self._read_pypdf2(path)
        except ImportError:
            raise ImportError(
                "请安装 pdfplumber 或 PyPDF2: "
                "pip install pdfplumber PyPDF2"
            )

    def _read_pdfplumber(self, path: Path) -> Document:
        import pdfplumber

        paragraphs = []
        with pdfplumber.open(str(path)) as pdf:
            total_pages = len(pdf.pages)
            for page_idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        paragraphs.append({
                            "index": len(paragraphs),
                            "text": line,
                            "style": f"page_{page_idx + 1}",
                        })

        doc = Document.from_paragraphs(path, "pdf", paragraphs)
        doc.metadata["page_count"] = total_pages
        return doc

    def _read_pypdf2(self, path: Path) -> Document:
        from PyPDF2 import PdfReader as PyPdfReader

        reader = PyPdfReader(str(path))
        paragraphs = []
        total_pages = len(reader.pages)

        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    paragraphs.append({
                        "index": len(paragraphs),
                        "text": line,
                        "style": f"page_{page_idx + 1}",
                    })

        doc = Document.from_paragraphs(path, "pdf", paragraphs)
        doc.metadata["page_count"] = total_pages
        return doc


# ── .txt 读取器 ───────────────────────────────────────────

class TxtReader(BaseReader):
    """读取 .txt 文件"""

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def read(self, path: Path) -> Document:
        for enc in [self.encoding, "utf-8", "gbk", "latin-1"]:
            try:
                text = path.read_text(encoding=enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            # 所有编码都失败，用 binary 读并忽略错误
            text = path.read_bytes().decode("utf-8", errors="replace")

        paragraphs = []
        for i, line in enumerate(text.split("\n")):
            line = line.strip()
            if line:
                paragraphs.append({
                    "index": i,
                    "text": line,
                    "style": "Normal",
                })

        return Document.from_paragraphs(path, "txt", paragraphs)


# ── .md 读取器 ───────────────────────────────────────────

class MdReader(BaseReader):
    """读取 .md 文件"""

    def read(self, path: Path) -> Document:
        text = path.read_text(encoding="utf-8")
        lines = text.split("\n")
        paragraphs = []

        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped:
                paragraphs.append({
                    "index": i,
                    "text": line,  # 保留原始缩进
                    "style": self._detect_style(stripped),
                })

        return Document.from_paragraphs(path, "md", paragraphs)

    @staticmethod
    def _detect_style(line: str) -> str:
        """检测 Markdown 行样式"""
        if line.startswith("# "):
            return "Heading 1"
        elif line.startswith("## "):
            return "Heading 2"
        elif line.startswith("### "):
            return "Heading 3"
        elif line.startswith("- ") or line.startswith("* "):
            return "List Item"
        elif line.startswith("|"):
            return "Table Row"
        elif "```" in line:
            return "Code Block"
        elif line.startswith(">"):
            return "Blockquote"
        else:
            return "Normal"


# ── 读取器注册表 ──────────────────────────────────────────

_reader_registry: dict[str, type[BaseReader]] = {
    ".docx": DocxReader,
    ".pdf": PdfReader,
    ".txt": TxtReader,
    ".md": MdReader,
}


def get_reader(path: Path) -> BaseReader:
    """根据文件后缀获取对应的读取器实例

    Args:
        path: 文件路径

    Returns:
        读取器实例

    Raises:
        ValueError: 不支持的文件格式
    """
    suffix = path.suffix.lower()
    reader_cls = _reader_registry.get(suffix)
    if reader_cls is None:
        supported = ", ".join(_reader_registry)
        raise ValueError(
            f"不支持的文件格式 '{suffix}'，支持的格式: {supported}"
        )
    return reader_cls()


def read_document(path: Path) -> Document:
    """便捷函数：读取文档

    Args:
        path: 文件路径

    Returns:
        Document 实例
    """
    reader = get_reader(path)
    logger.info("正在读取文档: %s (格式: %s)", path, path.suffix)
    return reader.read(path)


# ── 写入器基类 ──────────────────────────────────────────────

class BaseWriter(ABC):
    """文档写入器抽象基类"""

    @abstractmethod
    def write(self, document: Document, output_path: Path) -> Path:
        """将处理结果写入文件

        Args:
            document: 处理后的文档对象
            output_path: 输出路径

        Returns:
            写入的文件路径
        """
        ...


class TxtWriter(BaseWriter):
    """写入 .txt 文件"""

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def write(self, document: Document, output_path: Path) -> Path:
        output_path = output_path.with_suffix(".txt")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(document.content, encoding=self.encoding)
        logger.info("已写入 TXT: %s", output_path)
        return output_path


class DocxWriter(BaseWriter):
    """写入 .docx 文件，保持原段落结构"""

    def write(self, document: Document, output_path: Path) -> Path:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        output_path = output_path.with_suffix(".docx")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = DocxDoc()

        for para in document.paragraphs:
            text = para.get("text", "")
            style_name = para.get("style", "Normal")

            p = doc.add_paragraph()

            # 尝试匹配样式
            try:
                p.style = doc.styles[style_name]
            except (KeyError, AttributeError):
                pass

            run = p.add_run(text)
            run.font.name = "Times New Roman"

            # 标题加粗
            if style_name and "Heading" in style_name:
                run.bold = True
                level = style_name.replace("Heading ", "")
                if level.isdigit():
                    p.style = doc.styles[f"Heading {level}"]

        doc.save(str(output_path))
        logger.info("已写入 DOCX: %s", output_path)
        return output_path


# ── 写入调度 ──────────────────────────────────────────────

_writer_registry: dict[str, type[BaseWriter]] = {
    ".txt": TxtWriter,
    ".docx": DocxWriter,
}


def get_writer(format_str: str) -> BaseWriter:
    """获取写入器

    Args:
        format_str: 目标格式 (txt/docx/same_as_input)

    Returns:
        写入器实例

    Raises:
        ValueError: 不支持的输出格式
    """
    if format_str == "same_as_input":
        # 根据文档格式运行时再决定
        raise ValueError("same_as_input 需要在写入时根据输入格式确定目标格式")

    writer_cls = _writer_registry.get(f".{format_str.lower()}")
    if writer_cls is None:
        supported = ", ".join(_writer_registry)
        raise ValueError(
            f"不支持的输出格式 '{format_str}'，支持的格式: {supported}"
        )
    return writer_cls()


def write_document(document: Document, output_path: Path,
                   target_format: Optional[str] = None) -> Path:
    """便捷函数：写入文档

    Args:
        document: 处理后的文档
        output_path: 输出路径
        target_format: 目标格式，None 则根据 document.format 决定

    Returns:
        写入的文件路径
    """
    fmt = target_format or document.format
    if fmt == "same_as_input":
        fmt = document.format

    writer = get_writer(fmt)
    logger.info("正在写入文档: %s (格式: %s)", output_path, fmt)
    return writer.write(document, output_path)
