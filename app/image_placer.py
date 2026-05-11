"""智能图片插入模块

分析文档内容，确定图片的最佳插入位置。
支持用户指定位置（如"第3段后面""实验结果后面"）和自动识别两种模式。
"""

from __future__ import annotations

import logging
import re
import tempfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# 支持的图片格式
SUPPORTED_IMAGE_FORMATS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}

# 默认图片宽度（英寸）
DEFAULT_IMAGE_WIDTH = 5.0


class ImagePlacer:
    """智能图片放置器。

    分析文档内容，确定图片的最佳插入位置，使用纯关键词匹配，
    不依赖外部 API。
    """

    # 中文停用词
    _STOPWORDS = frozenset({
        '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都',
        '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会',
        '着', '没有', '看', '好', '自己', '这', '他', '她', '它', '们',
        '把', '被', '让', '给', '为', '所', '以', '能', '及', '与',
        '但', '而', '或', '如', '若', '虽', '因', '故', '乃', '其',
        '该', '此', '每', '各', '另', '那个', '这个', '什么', '怎么',
        '如何', '为何', '哪里', '哪些', '为何', '因为', '所以', '然而',
        '但是', '如果', '虽然', '而且', '并且', '或者', '还是', '只是',
        '不过', '关于', '对于', '通过', '根据', '按照', '经过', '本着',
    })

    def __init__(self):
        self._temp_files: list[str] = []

    def analyze_document(self, docx_path: str) -> list[dict]:
        """分析文档结构，返回段落信息列表。

        Args:
            docx_path: .docx 文件路径

        Returns:
            段落信息列表，每项包含：
            - index: 段落索引
            - text: 段落文本（去首尾空白）
            - style: 段落样式名
            - heading_level: 标题层级（0=普通段落，1-9=标题）
        """
        from docx import Document as DocxDoc

        doc = DocxDoc(docx_path)
        paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"
            heading_level = self._get_heading_level(style_name)

            paragraphs.append({
                "index": i,
                "text": text,
                "style": style_name,
                "heading_level": heading_level,
            })

        logger.info("文档分析完成: 共 %d 个非空段落", len(paragraphs))
        return paragraphs

    # ── 内部工具 ──────────────────────────────────────────────

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """从样式名中提取标题层级。

        Args:
            style_name: 段落样式名

        Returns:
            标题层级，0 表示普通段落，1-9 表示标题层级
        """
        if not style_name:
            return 0
        match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
        return int(match.group(1)) if match else 0

    # ── 位置推荐 ──────────────────────────────────────────────

    def suggest_position(self, paragraphs: list[dict],
                         image_description: str = "",
                         image_filename: str = "") -> dict:
        """根据文档内容和图片描述，推荐插入位置。

        算法：
        1. 提取图片关键词（从文件名或用户描述）
        2. 在文档中搜索相关段落
        3. 如果有强相关段落 → 插入到该段落后面
        4. 如果是配图 → 插入到最近的标题段落后面
        5. 如果完全无关 → 插入到文档末尾

        Args:
            paragraphs: 文档段落列表（analyze_document 的返回值）
            image_description: 图片描述（用户输入的说明）
            image_filename: 图片文件名（可用于提取关键词）

        Returns:
            dict: {paragraph_index, reason}
        """
        if not paragraphs:
            return {"paragraph_index": 0, "reason": "文档为空"}

        keywords = self._extract_keywords(image_description, image_filename)

        if not keywords:
            return {
                "paragraph_index": len(paragraphs) - 1,
                "reason": "无法提取图片关键词，已放在文档末尾",
            }

        # 在文档中搜索相关段落
        best_idx, score, matched_kws = self._find_best_paragraph(
            paragraphs, keywords
        )

        if score > 0:
            para = paragraphs[best_idx]
            level_info = f"（标题 H{para['heading_level']}）" if para["heading_level"] > 0 else ""
            kw_text = "、".join(matched_kws[:5])

            return {
                "paragraph_index": best_idx,
                "reason": (
                    f"文档第 {best_idx+1} 段{level_info}内容与图片关键词"
                    f"（{kw_text}）匹配（得分 {score}），"
                    f"图片放在此段落后最合适"
                ),
            }

        # 没有关键词匹配 → 找最近的标题段落
        for i in range(len(paragraphs) - 1, -1, -1):
            if paragraphs[i]["heading_level"] > 0:
                snippet = paragraphs[i]["text"][:40]
                return {
                    "paragraph_index": i,
                    "reason": (
                        f"未找到内容匹配的段落，"
                        f"图片放在最近的标题段落后（第 {i+1} 段：{snippet}...）"
                    ),
                }

        # 回退到末尾
        return {
            "paragraph_index": len(paragraphs) - 1,
            "reason": "未找到合适的插入位置，已放在文档末尾",
        }

    def _extract_keywords(self, description: str, filename: str) -> list[str]:
        """从描述和文件名中提取关键词。

        Args:
            description: 用户输入的描述
            filename: 图片文件名

        Returns:
            去重后的关键词列表
        """
        keywords: set[str] = set()

        # ── 从用户描述中提取 ──────────────────────────────
        if description:
            # 提取中文词组（2-4 字，非停用词）
            chinese = re.findall(r'[一-鿿]{2,4}', description)
            for word in chinese:
                if word not in self._STOPWORDS:
                    keywords.add(word)
                    # 对 4 字词拆出 2 字子关键词
                    if len(word) == 4:
                        sub_a, sub_b = word[:2], word[2:]
                        if sub_a not in self._STOPWORDS:
                            keywords.add(sub_a)
                        if sub_b not in self._STOPWORDS:
                            keywords.add(sub_b)
                    elif len(word) == 3:
                        sub = word[:2]
                        if sub not in self._STOPWORDS:
                            keywords.add(sub)

            # 提取英文单词
            english = re.findall(r'\b[a-zA-Z]{2,}\b', description)
            keywords.update(w.lower() for w in english)

        # ── 从文件名中提取 ──────────────────────────────
        if filename:
            stem = Path(filename).stem
            parts = re.split(r'[_\-\s.]+', stem)
            for part in parts:
                if len(part) >= 2:
                    keywords.add(part.lower())

        return list(keywords)

    def _find_best_paragraph(
        self, paragraphs: list[dict], keywords: list[str]
    ) -> tuple[int, int, list[str]]:
        """在段落中查找关键词匹配度最高的位置。

        对每个段落计算关键词匹配分数：
        - 基础分：匹配到的关键词数量 × 10
        - 标题加权：H2 × 1.5，H3 × 1.3
        - 密度加分：匹配占比 × 20
        - 返回最高分段落的（索引, 分数, 匹配到的关键词列表）

        Args:
            paragraphs: 段落列表
            keywords: 关键词列表

        Returns:
            (段落索引, 匹配分数, 匹配到的关键词列表)
        """
        if not paragraphs or not keywords:
            return (0, 0, [])

        best_idx = 0
        best_score = 0
        best_matched: list[str] = []

        for i, para in enumerate(paragraphs):
            text = para["text"].lower()
            heading_level = para["heading_level"]

            # 统计关键词匹配
            matched = [kw for kw in keywords if kw.lower() in text]
            match_count = len(matched)
            if match_count == 0:
                continue

            # 标题加权
            weight = 1.0
            if heading_level == 2:
                weight = 1.5
            elif heading_level == 3:
                weight = 1.3
            elif heading_level == 1:
                weight = 1.2

            score = int(match_count * weight * 10)
            # 密度加分
            density = int((match_count / len(keywords)) * 20)
            score += density

            if score > best_score:
                best_score = score
                best_idx = i
                best_matched = matched

        return (best_idx, best_score, best_matched)

    # ── 用户指令解析 ──────────────────────────────────────────

    def parse_user_instruction(self, instruction: str) -> dict:
        """解析用户的位置指令。

        规则：
        - 包含"第"+"段" + 数字 → 指定段落号
        - 包含"开头"/"开始" → 文档开头
        - 包含"末尾"/"结尾"/"最后" → 文档末尾
        - "X后面/X之后/…" → 关键词匹配
        - 其他 → 当作关键词

        Args:
            instruction: 用户输入的位置指令

        Returns:
            {type, value, position}
            type: "number" | "keyword" | "start" | "end" | "auto"
        """
        if not instruction or not instruction.strip():
            return {"type": "auto", "value": ""}

        text = instruction.strip()

        # 第X段 → 指定段落号
        m = re.search(r'第\s*(\d+)\s*段(?:落)?', text)
        if m:
            return {"type": "number", "value": int(m.group(1))}

        # 开头/开始
        if re.search(r'开头|开始|最前面|起始|文档头', text):
            return {"type": "start", "value": 0}

        # 末尾/结尾
        if re.search(r'末尾|结尾|最后|最后面|文档尾|文档末尾', text):
            return {"type": "end", "value": ""}

        # X后面/X之前/…  → 关键词 + 相对位置
        m = re.search(
            r'([一-鿿\w]{2,})\s*(后面|之后|下面|下方|'
            r'附近|旁边|前面|之前|上方|上面)',
            text,
        )
        if m:
            return {
                "type": "keyword",
                "value": m.group(1),
                "position": m.group(2),
            }

        # 包含"第" + 数字
        m = re.search(r'第\s*(\d+)\s*个', text)
        if m:
            return {"type": "number", "value": int(m.group(1))}

        # 其余统统当作关键词
        return {"type": "keyword", "value": text}

    # ── 图片处理 ──────────────────────────────────────────────

    def prepare_image(self, image_path: str, max_width: int = 500) -> str:
        """处理图片：调整大小、转换 RGB、压缩。

        Args:
            image_path: 原始图片路径
            max_width: 最大宽度（像素）

        Returns:
            处理后的图片路径（临时文件，用完需清理）
        """
        from PIL import Image

        img = Image.open(image_path)

        # RGBA → RGB（白色背景）
        if img.mode == 'RGBA':
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')

        # 按比例缩小
        if img.width > max_width:
            ratio = max_width / img.width
            new_h = int(img.height * ratio)
            img = img.resize((max_width, new_h), Image.LANCZOS)

        # 写入临时文件
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        tmp.close()
        img.save(tmp.name, 'PNG', optimize=True)
        self._temp_files.append(tmp.name)

        return tmp.name

    # ── 核心插入 ──────────────────────────────────────────────

    def place_image(self, docx_path: str, image_path: str,
                    position: Optional[dict] = None,
                    user_instruction: str = "") -> str:
        """在文档中插入图片。

        支持三种位置指定方式：
        1. position 参数：{paragraph_index, placement}
        2. user_instruction：自然语言指令（"第3段后面""实验结果后面"）
        3. 两者都为空：自动识别

        Args:
            docx_path: 原始 .docx 路径
            image_path: 图片文件路径
            position: {paragraph_index, placement("after"/"before"/"replace")}
            user_instruction: 用户描述

        Returns:
            修改后的 .docx 保存路径

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 参数无效
        """
        docx_path = Path(docx_path)
        img_path = Path(image_path)

        if not docx_path.exists():
            raise FileNotFoundError(f"文档不存在: {docx_path}")
        if not img_path.exists():
            raise FileNotFoundError(f"图片不存在: {image_path}")

        ext = img_path.suffix.lower()
        if ext not in SUPPORTED_IMAGE_FORMATS:
            raise ValueError(
                f"不支持的图片格式 '{ext}'，支持: "
                + ", ".join(SUPPORTED_IMAGE_FORMATS)
            )

        # ── 分析文档 ──────────────────────────────────────
        paragraphs = self.analyze_document(str(docx_path))
        total = len(paragraphs)
        if total == 0:
            raise ValueError("文档为空，无法插入图片")

        # ── 确定目标段落索引 ──────────────────────────────
        target_idx = self._resolve_target_index(
            paragraphs, position, user_instruction,
            Path(image_path).stem,
        )

        # ── 处理图片 ──────────────────────────────────────
        processed = self.prepare_image(str(image_path))

        # ── 插入图片 ──────────────────────────────────────
        from docx import Document as DocxDoc

        doc = DocxDoc(str(docx_path))
        para_count = len(doc.paragraphs)
        target_idx = max(0, min(target_idx, para_count - 1))

        placement = (position or {}).get("placement", "after")
        self._insert_image_at_paragraph(
            doc, target_idx, processed, placement=placement,
        )

        # ── 输出 ──────────────────────────────────────────
        output_path = docx_path.parent / f"{docx_path.stem}_插图版.docx"
        doc.save(str(output_path))
        logger.info("图片已插入: %s → %s", image_path, output_path)

        return str(output_path)

    def _resolve_target_index(self, paragraphs: list[dict],
                               position: Optional[dict],
                               user_instruction: str,
                               image_stem: str) -> int:
        """解析目标段落索引（三种方式优先级：position > 指令 > 自动）。"""
        total = len(paragraphs)

        if position and "paragraph_index" in position:
            idx = position["paragraph_index"]
            return max(0, min(idx, total - 1))

        if user_instruction:
            parsed = self.parse_user_instruction(user_instruction)

            if parsed["type"] == "number":
                return max(0, min(parsed["value"] - 1, total - 1))

            if parsed["type"] == "start":
                return 0

            if parsed["type"] == "end":
                return total - 1

            if parsed["type"] == "keyword":
                kws = self._extract_keywords(parsed["value"], "")
                if kws:
                    idx, score, _ = self._find_best_paragraph(paragraphs, kws)
                    if score > 0:
                        return idx
                # 关键词无匹配 → 末尾
                return total - 1

        # 自动识别
        suggestion = self.suggest_position(paragraphs, "", image_stem)
        return suggestion["paragraph_index"]

    @staticmethod
    def _insert_image_at_paragraph(doc, paragraph_index: int,
                                    image_path: str,
                                    width: float = DEFAULT_IMAGE_WIDTH,
                                    placement: str = "after"):
        """在指定位置插入包含图片的新段落。

        用 XML 树操作将新段落插入到目标段落的前/后，
        或替换目标段落内容。

        Args:
            doc: python-docx Document
            paragraph_index: 段落索引
            image_path: 已处理好的图片路径
            width: 图片宽度（英寸）
            placement: "after"（后）| "before"（前）| "replace"（替换）
        """
        from docx.shared import Inches
        from lxml import etree

        # 创建新段落并添加图片
        para = doc.paragraphs[paragraph_index]
        new_para = doc.add_paragraph()
        run = new_para.add_run()
        run.add_picture(image_path, width=Inches(width))

        if placement == "replace":
            # 替换：清空原段落文本，把图片 run 移过去
            for r in para.runs:
                r._element.getparent().remove(r._element)
            para._element.append(run._element)
            new_para._element.getparent().remove(new_para._element)

        elif placement == "before":
            # 前插
            para._element.addprevious(new_para._element)

        else:  # "after"
            # 后插
            para._element.addnext(new_para._element)

    # ── 清理 ──────────────────────────────────────────────────

    def cleanup(self):
        """清理当前实例产生的临时文件。"""
        for path_str in self._temp_files:
            try:
                Path(path_str).unlink(missing_ok=True)
            except OSError as e:
                logger.warning("清理临时文件失败: %s — %s", path_str, e)
        self._temp_files.clear()

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.cleanup()
