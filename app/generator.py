"""文档生成器

根据用户需求描述，调用 LLM 生成文档内容，保存为 .docx。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Callable, Optional

from .config import AppConfig
from .llm_client import create_client

logger = logging.getLogger(__name__)

# ── 默认提示词 ──────────────────────────────────────────────

DEFAULT_SYSTEM_PROMPT = """你是一个专业的文档写作助手。请根据用户的需求描述，生成一份结构完整、内容详实的文档。

要求：
1. 文档要有清晰的章节结构（标题层级）
2. 内容专业、具体、可操作
3. 如果提供了参考文档，请参考其风格和结构
4. 使用正式书面语
5. 适当使用表格、列表等格式"""

DEFAULT_USER_PROMPT_TEMPLATE = """请根据以下需求生成文档：

{requirement}

{reference_section}

文档格式要求：{format}格式"""

# ── 文档格式配置 ───────────────────────────────────────────

DOC_FORMATS = {
    '技术报告': {
        'description': '正式的技术报告风格',
        'suggested_sections': ['项目背景', '技术方案', '实施计划', '预算估算'],
    },
    '商业计划书': {
        'description': '商业计划书风格',
        'suggested_sections': ['项目概述', '市场分析', '商业模式', '财务预测'],
    },
    '项目方案': {
        'description': '项目实施方案风格',
        'suggested_sections': ['项目背景', '目标范围', '技术路线', '进度安排'],
    },
    '调研报告': {
        'description': '调研/研究报告风格',
        'suggested_sections': ['调研背景', '现状分析', '问题发现', '建议对策'],
    },
    '会议纪要': {
        'description': '会议纪要风格',
        'suggested_sections': ['会议信息', '讨论内容', '决议事项', '后续行动'],
    },
    '工作总结': {
        'description': '工作总结/汇报风格',
        'suggested_sections': ['工作概述', '完成情况', '问题分析', '下阶段计划'],
    },
    '自定义': {
        'description': '自由格式，按需求描述生成',
        'suggested_sections': [],
    },
}


class DocumentGenerator:
    """文档生成器。

    根据用户需求描述，调用 LLM 生成文档内容，保存为 .docx。
    """

    def __init__(self, config: AppConfig):
        self.config = config
        self._llm_client = create_client(config.llm)

    async def generate(
        self,
        requirement: str,
        reference_texts: list[str] = None,
        format: str = '技术报告',
        on_chunk: Callable[[str], None] = None,
    ) -> str:
        """生成文档内容。

        Args:
            requirement: 用户需求描述
            reference_texts: 参考文档文本（可选）
            format: 文档格式/类型
            on_chunk: 流式输出回调

        Returns:
            生成的完整文本
        """
        system_prompt, user_prompt = self._build_prompt(
            requirement, reference_texts or [], format
        )

        logger.info(
            "开始生成文档: format=%s, requirement_len=%d",
            format, len(requirement)
        )

        result = await self._llm_client.process_content(
            content="",
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            on_chunk=on_chunk,
            retry_count=self.config.processing.retry_count,
        )

        logger.info("文档生成完成: 共 %d 字符", len(result))
        return result

    def save_as_docx(
        self,
        text: str,
        output_path: str | Path,
        title: str = '未命名文档'
    ) -> Path:
        """将生成的文本保存为 .docx

        将 Markdown 风格的标题（# / ## / ###）映射为 Word 标题样式，
        识别表格（| 分隔行）和列表项（- / * / 数字），普通文本保持正文。

        Args:
            text: 生成的文本内容
            output_path: 输出路径
            title: 文档标题

        Returns:
            保存的文件路径
        """
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = DocxDoc()

        # ── 设置默认样式 ──────────────────────────────────
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 标题样式
        for level in range(1, 4):
            try:
                hs = doc.styles[f'Heading {level}']
                hs.font.name = 'Times New Roman'
                hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                hs.font.color.rgb = RGBColor(0, 0, 0)
                hs.font.bold = True
                hs.font.size = {1: Pt(18), 2: Pt(14), 3: Pt(12)}.get(level, Pt(12))
            except (KeyError, AttributeError):
                pass

        # ── 文档标题居中 ──────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(22)
        title_run.bold = True
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # ── 逐行解析内容 ──────────────────────────────────
        lines = text.split('\n')
        i = 0
        in_table = False
        table_rows: list[list[str]] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                # 空行：刷新未闭合表格
                if in_table and table_rows:
                    self._add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                i += 1
                continue

            # 标题：### 或 ## 或 #
            heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
            if heading_match:
                if in_table and table_rows:
                    self._add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                level = len(heading_match.group(1))
                doc.add_heading(heading_match.group(2), level=level)
                i += 1
                continue

            # 表格行：| a | b | c |
            if stripped.startswith('|') and stripped.endswith('|'):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if not in_table:
                    in_table = True
                    table_rows = [cells]
                else:
                    table_rows.append(cells)
                i += 1
                continue
            else:
                if in_table and table_rows:
                    self._add_table(doc, table_rows)
                    table_rows = []
                    in_table = False

            # 无序列表：- 或 * 开头
            list_match = re.match(r'^[-*+]\s+(.+)$', stripped)
            if list_match:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(list_match.group(1))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                i += 1
                continue

            # 有序列表：数字开头
            numbered_match = re.match(r'^\d+[.、]\s+(.+)$', stripped)
            if numbered_match:
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(numbered_match.group(1))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                i += 1
                continue

            # 普通正文
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1

        # 文件末尾未闭合表格
        if in_table and table_rows:
            self._add_table(doc, table_rows)

        doc.save(str(output_path))
        logger.info("DOCX 已保存: %s", output_path)
        return output_path

    # ── 内部方法 ──────────────────────────────────────────────

    def _add_table(self, doc, rows: list[list[str]]):
        """向文档中添加表格（首行加粗为表头）"""
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        if len(rows) < 2:
            return

        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= cols:
                    break
                cell = table.cell(row_idx, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if row_idx == 0:
                    run.bold = True

    def _build_prompt(
        self,
        requirement: str,
        reference_texts: list[str],
        format: str,
    ) -> tuple[str, str]:
        """构建 LLM 提示词

        Args:
            requirement: 用户需求描述
            reference_texts: 参考文档文本列表
            format: 文档格式

        Returns:
            (system_prompt, user_prompt)
        """
        # 参考文档部分
        reference_section = ""
        if reference_texts:
            ref_blocks = []
            for idx, ref_text in enumerate(reference_texts):
                # 截取前 3000 字符避免超长
                truncated = ref_text[:3000]
                if len(ref_text) > 3000:
                    truncated += "\n...（剩余内容已截断）"
                ref_blocks.append(
                    f"--- 参考文档 {idx + 1} ---\n{truncated}"
                )
            reference_section = (
                "以下是参考文档的内容，请参考其风格和结构：\n\n"
                + "\n\n".join(ref_blocks)
            )

        # 格式说明
        format_info = DOC_FORMATS.get(format, DOC_FORMATS['自定义'])
        format_desc = format_info['description']
        if format_info['suggested_sections']:
            format_desc += (
                f"\n建议包含以下章节："
                f"{'、'.join(format_info['suggested_sections'])}"
            )

        user_prompt = DEFAULT_USER_PROMPT_TEMPLATE.format(
            requirement=requirement,
            reference_section=reference_section,
            format=format,
        )
        user_prompt += f"\n\n格式说明：{format_desc}"

        from app.processing_skill import ENCODING_SAFETY_PROMPT
        return DEFAULT_SYSTEM_PROMPT + '\n\n' + ENCODING_SAFETY_PROMPT, user_prompt
