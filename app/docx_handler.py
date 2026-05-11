"""DOCX 文件处理器

读取 .docx 文件段落和格式信息，AI 处理文本后写回文档，
尽量保留原格式（字体、字号、加粗/斜体、对齐方式、缩进等）。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocxHandler:
    """DOCX 文件处理器

    使用方式：
        handler = DocxHandler()
        handler.open('input.docx')
        text = handler.get_text()       # 给 AI 处理
        handler.apply_changes(new_text)  # AI 结果写回
        handler.save('output.docx')
    """

    def __init__(self):
        self._doc = None
        self._file_path = None
        # 存储每段原始格式信息
        self._para_formats: list[dict] = []

    # ── 公开接口 ─────────────────────────────────────────────

    def open(self, file_path: str | Path):
        """打开 .docx 文件，解析段落和格式。

        Args:
            file_path: .docx 文件路径
        """
        from docx import Document as DocxDoc

        self._file_path = Path(file_path)
        self._doc = DocxDoc(str(self._file_path))
        self._parse_formats()

    def get_text(self) -> str:
        """提取纯文本（用于 AI 处理）。

        Returns:
            文档全部文本，段落间用换行分隔
        """
        if self._doc is None:
            return ""
        return "\n".join(
            p.text for p in self._doc.paragraphs
        )

    def get_paragraphs(self) -> list[dict]:
        """获取段落列表，每段包含格式信息。

        Returns:
            [{text, style, font_name, font_size, bold, italic, alignment, ...}]
        """
        return list(self._para_formats)

    def apply_changes(self, new_text: str):
        """将 AI 处理后的文本写回文档，尽量保留原格式。

        策略：
          - 如果新段落数与原段落数相同，逐段保留格式。
          - 如果新段落数不同，按比例映射（或复用最后一段的格式）。
          - 表格、图片等非段落元素保持不动。

        Args:
            new_text: AI 处理后的文本，段落间用换行分隔
        """
        if self._doc is None:
            raise RuntimeError("尚未打开文档，请先调用 open()")

        new_paras = new_text.split("\n")
        old_count = len(self._para_formats)
        new_count = len(new_paras)

        # 清除原有段落文本，保留非段落内容（表格等）
        self._clear_paragraph_texts()

        # 按格式映射填充新文本
        for i, para_text in enumerate(new_paras):
            fmt = self._map_format(i, old_count, new_count)
            if i < old_count:
                # 重用原有段落
                self._apply_format_to_paragraph(
                    self._doc.paragraphs[i], para_text, fmt
                )
            else:
                # 超出原段数，追加新段落
                p = self._doc.add_paragraph()
                self._apply_format_to_paragraph(p, para_text, fmt)

    def save(self, output_path: str | Path):
        """保存修改后的 .docx。

        Args:
            output_path: 输出文件路径
        """
        if self._doc is None:
            raise RuntimeError("尚未打开文档，请先调用 open()")
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        # 设置中国国家标准页面边距（上2.54cm、下2.54cm、左3.17cm、右3.17cm）
        from docx.shared import Cm
        for section in self._doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._doc.save(str(output_path))
        logger.info("已保存 DOCX: %s", output_path)

    # ── 私有方法 ─────────────────────────────────────────────

    def _parse_formats(self):
        """遍历段落，提取每段的格式信息。"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._para_formats = []
        for para in self._doc.paragraphs:
            fmt = {
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "alignment": para.alignment,
                "first_line_indent": para.paragraph_format.first_line_indent,
                "left_indent": para.paragraph_format.left_indent,
                "space_before": para.paragraph_format.space_before,
                "space_after": para.paragraph_format.space_after,
                "line_spacing": para.paragraph_format.line_spacing,
                "runs": [],
            }
            # 提取每个 run 的格式
            for run in para.runs:
                run_info = {
                    "text": run.text,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "color": (
                        run.font.color.rgb
                        if run.font.color and run.font.color.rgb
                        else None
                    ),
                }
                fmt["runs"].append(run_info)

            self._para_formats.append(fmt)

    def _clear_paragraph_texts(self):
        """清除所有段落文本（保留 run 框架以便后续恢复格式）。

        只清除纯文本 run，不触碰包含图片（w:drawing / w:pict）的 run，
        避免丢失文档中的图片关系引用。
        """
        from docx.oxml.ns import qn

        for para in self._doc.paragraphs:
            for run in para.runs:
                # 跳过包含图片的 run（drawing 或 pict 元素）
                has_image = (
                    run._element.findall(qn('w:drawing'))
                    or run._element.findall(qn('w:pict'))
                )
                if has_image:
                    continue
                run.text = ""

    def _map_format(self, idx: int, old_count: int, new_count: int) -> dict:
        """根据段落索引映射到对应的原始格式。

        Args:
            idx: 新段落索引
            old_count: 原段落数
            new_count: 新段落数

        Returns:
            对应的原始段落格式
        """
        if old_count == 0:
            return {
                "style": "Normal",
                "alignment": None,
                "first_line_indent": None,
                "left_indent": None,
                "space_before": None,
                "space_after": None,
                "runs": [],
            }

        if old_count == new_count:
            # 段落数相同，直接对应
            orig_idx = idx
        else:
            # 按比例映射
            ratio = idx / max(new_count - 1, 1)
            orig_idx = min(int(ratio * old_count), old_count - 1)

        orig_idx = min(orig_idx, old_count - 1)
        return self._para_formats[orig_idx]

    def _apply_format_to_paragraph(self, paragraph, text: str, fmt: dict):
        """将格式应用到段落，按中国国家标准学术格式设置。

        Args:
            paragraph: python-docx Paragraph 对象
            text: 段落文本
            fmt: 格式字典（来自 _para_formats）
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # 不要使用 paragraph.clear() — 它会删除包含图片的 run（w:drawing/w:pict），
        # 导致文档中的图片和对应关系引用丢失。
        # 改为仅移除纯文本 run，保留图片 run。
        runs_to_remove = [
            r for r in paragraph.runs
            if not (
                r._element.findall(qn('w:drawing'))
                or r._element.findall(qn('w:pict'))
            )
        ]
        for r in runs_to_remove:
            r._element.getparent().remove(r._element)

        # 保留空段落（纯空格、空文本），不添加空 run
        if not text:
            return

        run = paragraph.add_run(text)
        style_name = str(fmt.get("style", "")).lower()
        original_runs = fmt.get("runs")
        is_heading = 'heading' in style_name or any(k in style_name for k in ['标题', 'head'])

        # 非标题段落：优先保留原始格式
        if not is_heading and original_runs:
            first = original_runs[0]
            fn = first.get("font_name")
            fs = first.get("font_size")
            b = first.get("bold")
            if fn:
                run.font.name = fn
                run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
            if fs:
                run.font.size = fs
            if b is not None:
                run.bold = b
            orig_line = fmt.get('line_spacing')
            if orig_line:
                paragraph.paragraph_format.line_spacing = orig_line
            return

        # 中文论文标准格式（标题段落或无线索时应用）
        if is_heading:
            level = 1
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.bold = True
            sizes = {1: 15, 2: 14, 3: 12}
            run.font.size = Pt(sizes.get(level, 14))
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
        elif 'normal' in style_name or style_name in ('', '正文'):
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.line_spacing = 1.5
        else:
            # 默认正文格式
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(24)

        # 英文/数字用 Times New Roman
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def _fill_runs(self, paragraph, text: str, runs: list[dict]):
        """将文本按比例填入多个 run，保留每个 run 的格式。"""
        total_run_chars = sum(len(r["text"]) for r in runs) or 1
        start = 0

        for run_info in runs:
            # 计算本 run 应分配多少字符
            run_len_ratio = len(run_info["text"]) / total_run_chars
            end = start + max(1, int(len(text) * run_len_ratio))
            if end > len(text):
                end = len(text)

            run_text = text[start:end] if start < len(text) else ""
            run = paragraph.add_run(run_text)

            # 还原格式
            if run_info.get("font_name"):
                run.font.name = run_info["font_name"]
            if run_info.get("font_size"):
                run.font.size = run_info["font_size"]
            if run_info.get("bold") is not None:
                run.bold = run_info["bold"]
            if run_info.get("italic") is not None:
                run.italic = run_info["italic"]
            if run_info.get("underline") is not None:
                run.underline = run_info["underline"]
            if run_info.get("color"):
                from docx.shared import RGBColor
                run.font.color.rgb = run_info["color"]

            start = end
