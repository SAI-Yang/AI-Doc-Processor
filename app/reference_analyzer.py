"""参考文档分析器

纯规则引擎，分析参考文档的结构、风格、格式、关键词。
不调用 LLM，所有分析基于统计和规则匹配。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── 中文停用词 ──────────────────────────────────────────────
_CHINESE_STOPWORDS: set[str] = {
    "的", "了", "在", "是", "我", "有", "和", "就", "不", "人",
    "都", "一", "一个", "上", "也", "很", "到", "说", "要", "去",
    "你", "会", "着", "没有", "看", "好", "自己", "这", "他", "她",
    "它", "们", "那", "些", "来", "为", "与", "及", "等", "被",
    "把", "对", "从", "以", "之", "而", "但", "或", "于", "因",
    "所", "如", "更", "将", "并", "其", "中", "者", "各", "让",
    "可以", "这个", "那个", "什么", "因为", "所以", "如果", "虽然",
    "但是", "而且", "然后", "已经", "可能", "应该", "能够", "通过",
    "进行", "使用", "采用", "利用", "基于", "关于", "按照", "根据",
    "需要", "能够", "可以", "没有", "不是", "就是", "只是", "但是",
    "还是", "或是", "或是", "属于", "成为", "作为", "认为", "称为",
    "以及", "及其", "之中", "之一", "之间", "以后", "以来", "以外",
    "之前", "之后", "以上", "以下", "左右", "上下", "多个", "各种",
    "不同", "主要", "重要", "基本", "一般", "特别", "非常", "比较",
    "方面", "方式", "方法", "情况", "问题", "内容", "部分", "类型",
    "结果", "过程", "信息", "数据", "技术", "系统", "相关", "这种",
    "这个", "这些", "那些", "其他", "整个", "同时", "当时", "目前",
    "当前", "现在", "未来", "过去", "原有", "新的", "现有", "所有",
}

# ── 正式度关键词 ────────────────────────────────────────────
_FORMAL_WORDS: set[str] = {
    "笔者", "本文", "本研究", "综上所述", "因此", "然而",
    "鉴于", "据此", "由此可见", "亦即", "换言之", "诚然",
    "尚且", "固然", "未免", "颇为", "极为", "较为",
    "显著", "明显", "充分", "相应", "相关", "上述", "如下",
    "establish", "demonstrate", "consequently", "furthermore",
    "moreover", "nevertheless", "notwithstanding", "subsequently",
    "therefore", "thus", "whereas", "hence", "thereby",
}

_CASUAL_WORDS: set[str] = {
    "我们", "我觉得", "我认为", "个人认为", "大概", "可能吧",
    "好像", "感觉", "有点", "挺", "就这样", "嗯", "哦",
    "好吧", "对了", "其实", "反正", "话说", "说白了",
    "we think", "i think", "i feel", "maybe", "sort of",
    "kind of", "pretty good", "anyway", "well",
    "you know", "by the way",
}

# ── 技术术语模式 ────────────────────────────────────────────
_TECH_TERM_PATTERNS: list[re.Pattern] = [
    re.compile(r'\b[A-Z]{2,}(?:s|es)?\b'),                            # 全大写缩写: API, DSP, PDF
    re.compile(r'\b[A-Z][a-z]+[A-Z]\w*\b'),                           # 驼峰式: DeepSeek, PyQt5
    re.compile(r'\b\d+(?:\.\d+)?\s*(?:Hz|MHz|GHz|KB|MB|GB|TB|dpi|mm|cm|nm|ms|us)\b', re.IGNORECASE),  # 数值+单位
    re.compile(r'\b[a-z]+[\d][a-z0-9]*\b', re.IGNORECASE),           # 含数字: STM32, FFT512
]

# ── 主观/说服性语气词 ───────────────────────────────────────
_SUBJECTIVE_WORDS: list[str] = [
    "我认为", "我觉得", "在我看来", "我个人",
    "i think", "in my opinion", "i believe", "personally",
    "可能", "大概", "maybe", "perhaps", "probably",
]

_PERSUASIVE_WORDS: list[str] = [
    "必须", "应该", "务必", "一定要", "毫无疑问",
    "显然", "必然", "绝对", "肯定",
    "must", "should", "must be", "undoubtedly",
    "clearly", "obviously", "absolutely", "certainly",
]


def _detect_language(text: str) -> str:
    """检测主要语言

    统计中英文字符数，中文多返回 "zh"，英文多返回 "en"。

    Args:
        text: 待检测文本

    Returns:
        "zh" 或 "en"
    """
    if not text.strip():
        return "zh"
    chinese = len(re.findall(r'[一-鿿]', text))
    english = len(re.findall(r'[a-zA-Z]', text))
    return "zh" if chinese >= english else "en"


def _extract_text_from_file(file_path: Path) -> str:
    """从文件中提取纯文本内容

    Args:
        file_path: 文件路径

    Returns:
        文件纯文本内容
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        from docx import Document as DocxDoc
        doc = DocxDoc(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)

    elif suffix == ".pdf":
        # 优先 pdfplumber，回退 PyPDF2
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                return "\n".join(pages)
        except ImportError:
            pass

        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(file_path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raise ImportError("请安装 pdfplumber 或 PyPDF2 以支持 PDF 分析")

    else:
        # txt, md 等纯文本
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return file_path.read_text(encoding=enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return file_path.read_bytes().decode("utf-8", errors="replace")


def _extract_keywords_tf(text: str, language: str, top_n: int = 10) -> list[str]:
    """基于词频的关键词提取（纯规则）

    Args:
        text: 文档文本
        language: "zh" 或 "en"
        top_n: 返回前 N 个关键词

    Returns:
        关键词列表
    """
    if language == "zh":
        try:
            import jieba
            words = list(jieba.cut(text))
        except ImportError:
            # 无 jieba 时按单字 + 双字组合
            words = []
            for c in text:
                if '一' <= c <= '鿿':
                    words.append(c)
            # 简单双字组合
            chars = [c for c in text if '一' <= c <= '鿿']
            for i in range(len(chars) - 1):
                words.append(chars[i] + chars[i + 1])
    else:
        words = re.findall(r'[a-zA-Z][a-zA-Z\'-]*', text.lower())

    # 过滤停用词、短词、纯数字
    filtered: list[str] = []
    for w in words:
        w_clean = w.strip().lower()
        if len(w_clean) < 2:
            continue
        if w_clean.isdigit():
            continue
        if w_clean in _CHINESE_STOPWORDS:
            continue
        filtered.append(w_clean)

    # 词频统计
    freq: dict[str, int] = {}
    for w in filtered:
        freq[w] = freq.get(w, 0) + 1

    # 排序取 top_n
    sorted_words = sorted(freq.items(), key=lambda x: (-x[1], x[0]))
    return [w for w, _ in sorted_words[:top_n]]


def _analyze_docx_structure(file_path: Path) -> tuple[list[dict], int, int, int, int]:
    """分析 docx 文档的结构、表格和图片

    Args:
        file_path: .docx 文件路径

    Returns:
        (structure, heading_count, paragraph_count, table_count, image_count)
    """
    from docx import Document as DocxDoc

    doc = DocxDoc(str(file_path))

    # 结构分析：识别标题层级
    structure: list[dict] = []
    stack: list[list | dict] = [structure]
    heading_count = 0
    para_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_count += 1

        style_name = para.style.name if para.style else ""

        level = 0
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except (ValueError, AttributeError):
                level = 1
            heading_count += 1

        if level > 0:
            node: dict = {"level": level, "text": text, "children": []}

            # 弹出栈顶直到找到合适层级的父级
            # stack 结构: [structure, node_lv1, node_lv2, ...]
            while len(stack) > level:
                stack.pop()
            # 如果栈太浅，补足到目标层级
            while len(stack) <= level:
                parent = stack[-1]
                if isinstance(parent, list):
                    parent.append(node)
                else:
                    parent["children"].append(node)
                stack.append(node)
                break

    # 表格统计
    table_count = len(doc.tables)

    # 图片统计（通过关系）
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    return structure, heading_count, para_count, table_count, image_count


def _analyze_md_structure(text: str) -> tuple[list[dict], int, int]:
    """分析 Markdown 文档结构

    Args:
        text: Markdown 文本

    Returns:
        (structure, heading_count, paragraph_count)
    """
    structure: list[dict] = []
    heading_count = 0
    para_count = 0

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        para_count += 1

        if stripped.startswith("# "):
            structure.append({"level": 1, "text": stripped[2:].strip(), "children": []})
            heading_count += 1
        elif stripped.startswith("## "):
            structure.append({"level": 2, "text": stripped[3:].strip(), "children": []})
            heading_count += 1
        elif stripped.startswith("### "):
            structure.append({"level": 3, "text": stripped[4:].strip(), "children": []})
            heading_count += 1

    return structure, heading_count, para_count


# ── 分析器类 ────────────────────────────────────────────────


class ReferenceAnalyzer:
    """参考文档分析器

    分析参考文档的结构、风格、格式，提取关键信息供生成/修改使用。
    所有分析基于规则引擎，不调用 LLM。
    """

    def analyze(self, file_path: str) -> dict:
        """分析参考文档，返回结构化信息

        Args:
            file_path: 文件路径

        Returns:
            分析结果字典，包含 title, language, word_count, paragraph_count,
            structure, style, key_topics, summary, table_count, image_count

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in (".docx", ".pdf", ".txt", ".md"):
            raise ValueError(f"不支持的文件格式: {suffix}")

        # 提取文本
        text = _extract_text_from_file(path)

        # 结构分析（格式相关）
        if suffix == ".docx":
            structure, heading_count, para_count, table_count, image_count = \
                _analyze_docx_structure(path)
        elif suffix == ".md":
            structure, heading_count, para_count = _analyze_md_structure(text)
            table_count = 0
            image_count = 0
        else:
            # pdf, txt — 无法精确识别标题结构
            structure = []
            heading_count = 0
            para_count = len([l for l in text.split("\n") if l.strip()])
            table_count = 0
            image_count = 0

        # 语言检测
        language = _detect_language(text)

        # 风格分析
        style = self._analyze_style(text, path)
        style["has_tables"] = table_count > 0
        style["has_images"] = image_count > 0

        # 关键词提取
        key_topics = _extract_keywords_tf(text, language)

        # 摘要（前 500 字）
        summary = text[:500].strip() if text else ""
        if len(text) > 500:
            summary += "..."

        # 字数统计
        if language == "zh":
            word_count = len(re.findall(r'[一-鿿]', text))
        else:
            word_count = len(text.split())

        return {
            "title": path.stem,
            "language": language,
            "word_count": word_count,
            "paragraph_count": para_count,
            "structure": structure,
            "heading_count": heading_count,
            "style": style,
            "key_topics": key_topics,
            "summary": summary,
            "table_count": table_count,
            "image_count": image_count,
        }

    # ── 风格分析 ──────────────────────────────────────────

    def _analyze_style(self, text: str, path: Path) -> dict:
        """分析文档风格

        Args:
            text: 文档文本
            path: 文件路径（用于提取字体等格式信息）

        Returns:
            风格字典
        """
        # 正式度
        formal_count = sum(1 for w in _FORMAL_WORDS if w in text.lower())
        casual_count = sum(1 for w in _CASUAL_WORDS if w in text.lower())

        if formal_count > casual_count * 2:
            formality = "formal"
        elif casual_count > formal_count * 2:
            formality = "casual"
        elif formal_count > 0:
            formality = "formal"
        else:
            formality = "casual"

        # 语气
        subj_count = sum(1 for w in _SUBJECTIVE_WORDS if w in text.lower())
        pers_count = sum(1 for w in _PERSUASIVE_WORDS if w in text.lower())

        if subj_count == 0 and pers_count == 0:
            tone = "objective"
        elif pers_count > subj_count:
            tone = "persuasive"
        elif subj_count > pers_count:
            tone = "subjective"
        else:
            tone = "objective"

        # 技术性 — 检测专业术语密度
        tech_count = 0
        for pattern in _TECH_TERM_PATTERNS:
            tech_count += len(pattern.findall(text))

        # 提取字体格式（仅 docx）
        font_name = ""
        font_size = ""
        line_spacing = 1.5

        if path.suffix.lower() == ".docx":
            try:
                from docx import Document as DocxDoc
                doc = DocxDoc(str(path))

                # 找第一个有格式的段落
                for para in doc.paragraphs:
                    if para.runs:
                        run = para.runs[0]
                        if run.font.name:
                            font_name = run.font.name
                        if run.font.size:
                            font_size = f"{run.font.size.pt:.0f}pt"
                        break

                # 行间距
                for para in doc.paragraphs:
                    ls = para.paragraph_format.line_spacing
                    if ls:
                        line_spacing = ls
                        break
            except Exception as exc:
                logger.debug("提取 docx 格式信息失败: %s", exc)

        return {
            "font": font_name,
            "font_size": font_size,
            "line_spacing": line_spacing,
            "has_tables": False,   # 由调用方覆盖
            "has_images": False,   # 由调用方覆盖
            "formality": formality,
            "tone": tone,
            "tech_term_count": tech_count,
        }

    # ── 样式模板提取 ──────────────────────────────────────

    def extract_style_template(self, docx_path: str) -> dict:
        """提取文档的样式模板（字体、字号、颜色、间距等）

        Args:
            docx_path: .docx 文件路径

        Returns:
            {
                "styles": {style_name: {font_name, font_size, bold, italic, ...}},
                "paragraph_counts": {style_name: count},
            }

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不是 docx 格式
        """
        from docx import Document as DocxDoc

        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        if path.suffix.lower() != ".docx":
            raise ValueError("仅支持 .docx 格式的样式提取")

        doc = DocxDoc(str(path))

        styles: dict[str, dict] = {}
        for para in doc.paragraphs:
            if not para.runs:
                continue
            style_name = para.style.name if para.style else "Normal"
            if style_name not in styles:
                run = para.runs[0]
                style_info: dict = {
                    "font_name": run.font.name or "",
                    "font_size": round(run.font.size.pt, 1) if run.font.size else 11.0,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else "",
                    "alignment": str(para.alignment) if para.alignment else "",
                    "first_line_indent": str(para.paragraph_format.first_line_indent) if para.paragraph_format.first_line_indent else "",
                    "line_spacing": para.paragraph_format.line_spacing or 1.5,
                    "space_before": str(para.paragraph_format.space_before) if para.paragraph_format.space_before else "",
                    "space_after": str(para.paragraph_format.space_after) if para.paragraph_format.space_after else "",
                }
                styles[style_name] = style_info

        # 各样式段落数统计
        para_counts: dict[str, int] = {}
        for para in doc.paragraphs:
            name = para.style.name if para.style else "Normal"
            para_counts[name] = para_counts.get(name, 0) + 1

        return {
            "styles": styles,
            "paragraph_counts": para_counts,
        }

    # ── 需求对比 ──────────────────────────────────────────

    def compare_with_requirement(self, analysis: dict, requirement: str) -> dict:
        """对比用户需求与参考文档分析结果，返回参考建议

        Args:
            analysis: analyze() 返回的分析结果
            requirement: 用户的处理需求（如"按这个风格来写"）

        Returns:
            {
                "language": "zh"/"en",
                "suggestions": [
                    {"type": "structure", "message": "...", "details": [...]},
                    ...
                ],
            }
        """
        suggestions: list[dict] = []
        req_lower = requirement.lower()

        # 结构参考建议
        heading_count = analysis.get("heading_count", 0)
        if heading_count > 0 and any(kw in requirement for kw in ["结构", "格式", "章节", "structure", "format"]):
            structure = analysis.get("structure", [])
            titles = [s.get("text", "")[:40] for s in structure[:8]]
            suggestions.append({
                "type": "structure",
                "message": f"参考文档包含 {heading_count} 个标题层级，建议参考其章节结构",
                "details": titles,
            })

        # 风格参考建议
        style = analysis.get("style", {})
        if any(kw in requirement for kw in ["风格", "样式", "style", "formal"]):
            detail_lines = []
            if style.get("font"):
                detail_lines.append(f"字体: {style['font']} {style.get('font_size', '')}")
            detail_lines.append(f"语气: {style.get('tone', '客观')}")
            suggestions.append({
                "type": "style",
                "message": f"参考文档风格为「{style.get('formality', '未知')}」，{style.get('tone', '客观')}语气",
                "details": detail_lines,
            })

        # 语言参考建议
        lang = analysis.get("language", "zh")
        if any(kw in requirement for kw in ["语言", "language", "中文", "英文"]):
            suggestions.append({
                "type": "language",
                "message": f"参考文档主要语言为 {'中文' if lang == 'zh' else '英文'}",
                "details": [],
            })

        # 关键词参考
        topics = analysis.get("key_topics", [])
        if topics:
            suggestions.append({
                "type": "topics",
                "message": f"参考文档关键词: {'、'.join(topics[:6])}",
                "details": topics[:10],
            })

        # 表格/图片参考
        if analysis.get("table_count", 0) > 0:
            suggestions.append({
                "type": "tables",
                "message": f"参考文档包含 {analysis['table_count']} 个表格",
                "details": [],
            })
        if analysis.get("image_count", 0) > 0:
            suggestions.append({
                "type": "images",
                "message": f"参考文档包含 {analysis['image_count']} 张图片",
                "details": [],
            })

        return {
            "language": lang,
            "suggestions": suggestions,
        }
