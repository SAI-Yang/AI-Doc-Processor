"""科学图表生成引擎

根据自然语言描述生成 Nature 期刊风格的科学图表。
使用 Matplotlib 绘制，Pillow 后处理。
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
import textwrap
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Nature 期刊调色板 ────────────────────────────────────────
NATURE_COLORS = [
    '#3B4992', '#EE0000', '#008B45', '#631879',
    '#008280', '#BB0021', '#5F559B', '#A20056',
]

# ── Matplotlib 中文字体探测 ──────────────────────────────────

# 常见支持中文的字体名称（按优先级排序）
_CJK_FONTS = [
    'Microsoft YaHei',       # Windows 雅黑
    'SimHei',                # Windows 黑体
    'PingFang SC',           # macOS
    'Noto Sans CJK SC',      # Linux
    'WenQuanYi Micro Hei',   # Linux
    'Arial Unicode MS',     # macOS fallback
]


def _find_cjk_font() -> str:
    """探测系统中可用的中文字体名，返回第一个匹配的字体。"""
    import matplotlib.font_manager as fm
    candidates = set()
    for f in fm.fontManager.ttflist:
        candidates.add(f.name)
    for name in _CJK_FONTS:
        if name in candidates:
            return name
    # fallback
    return 'sans-serif'


# 全局字体缓存
_CJK_FONT_NAME: str = ''


def _ensure_cjk_font() -> str:
    global _CJK_FONT_NAME
    if not _CJK_FONT_NAME:
        _CJK_FONT_NAME = _find_cjk_font()
    return _CJK_FONT_NAME


# ═══════════════════════════════════════════════════════════════
#  FigureGenerator
# ═══════════════════════════════════════════════════════════════

class FigureGenerator:
    """科学图表生成器

    根据描述生成 Nature 风格的科学图表。
    使用 Matplotlib 绘制，Pillow 后处理。

    Usage:
        gen = FigureGenerator()
        path = gen.generate("绘制x轴为频率、y轴为幅值的频谱图",
                            "output.png", chart_type='auto')
    """

    CHART_TYPES = {
        'line': '折线图',
        'bar': '柱状图',
        'scatter': '散点图',
        'histogram': '直方图',
        'boxplot': '箱线图',
        'heatmap': '热力图',
        'multiline': '多线对比图',
        'grouped_bar': '分组柱状图',
    }

    # 内置图表模板注册表
    _TEMPLATES: dict = {}

    def __init__(self, dpi: int = 300, figsize: tuple = (3.5, 2.5)):
        """
        Args:
            dpi: 输出分辨率
            figsize: 默认图形尺寸（英寸），Nature 单栏宽 3.5in
        """
        self._dpi = dpi
        self._figsize = figsize
        # 延迟导入 matplotlib（只在需要时加载）
        self._imported = False

    # ── 延迟导入 ─────────────────────────────────────────────

    def _lazy_import(self):
        """延迟导入 matplotlib/pyplot，避免阻塞 GUI 启动。"""
        if self._imported:
            return
        import matplotlib
        matplotlib.use('Agg')  # 非交互式后端
        import matplotlib.pyplot as plt
        import matplotlib.ticker as ticker
        from matplotlib import rcParams

        self._plt = plt
        self._ticker = ticker
        self._rcParams = rcParams

        # 全局样式 —— 接近 Nature 默认
        rcParams.update({
            'font.family': 'sans-serif',
            'font.sans-serif': [_ensure_cjk_font(), 'Arial', 'Helvetica'],
            'font.size': 7.0,
            'axes.linewidth': 0.5,
            'axes.labelsize': 8,
            'axes.titlesize': 8,
            'xtick.labelsize': 7,
            'ytick.labelsize': 7,
            'legend.fontsize': 7,
            'lines.linewidth': 1.0,
            'lines.markersize': 3.0,
            'figure.dpi': self._dpi,
            'savefig.dpi': self._dpi,
            'savefig.bbox': 'tight',
            'savefig.pad_inches': 0.05,
        })

        self._imported = True

    # ── Nature 风格应用 ──────────────────────────────────────

    def nature_style(self, ax) -> None:
        """应用 Nature 期刊风格的坐标轴样式。

        Nature 规范：
        - 白色背景，无网格
        - 坐标轴黑色，线宽 0.5pt
        - 字体 Arial/Helvetica，6-8pt
        - 刻度朝内
        - 无边框（右侧和上侧隐藏）
        - 颜色使用 Nature 调色板
        """
        # 隐藏上侧和右侧边框
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        # 剩余边框线宽
        for side in ('bottom', 'left'):
            ax.spines[side].set_linewidth(0.5)
            ax.spines[side].set_color('black')

        # 无网格
        ax.grid(False)

        # 刻度朝内
        ax.tick_params(
            axis='both', direction='in', length=2.5, width=0.5,
            colors='black', pad=3,
        )

        # 背景白色
        ax.set_facecolor('white')
        ax.patch.set_visible(False)

    def nature_colors(self, n: int) -> list[str]:
        """获取 Nature 调色板中的前 n 种颜色。

        Args:
            n: 需要的颜色数量

        Returns:
            颜色 HEX 字符串列表
        """
        colors = NATURE_COLORS[:]
        # 如果 n 大于调色板长度，重复轮转
        while len(colors) < n:
            colors.extend(NATURE_COLORS)
        return colors[:n]

    # ── 主入口 ───────────────────────────────────────────────

    def generate(self, description: str, output_path: str,
                 chart_type: str = 'auto',
                 data: Optional[dict] = None) -> str:
        """根据描述生成图表。

        流程：
        1. 分析描述，确定图表类型和数据形式
        2. 查找或自动生成数据
        3. 选择对应模板绘制
        4. 应用 Nature 风格
        5. 保存为高分辨率 PNG

        Args:
            description: 用户描述（如"绘制x轴为频率、y轴为幅值的频谱图"）
            output_path: 输出图片路径（应包含 .png 后缀）
            chart_type: 图表类型，'auto' 表示自动识别; 也可指定 CHART_TYPES 中的 key
            data: 可选的结构化数据，格式取决于图表类型

        Returns:
            生成的图片路径

        Raises:
            ValueError: 图表类型不合法或描述无法解析
            RuntimeError: 渲染失败
        """
        self._lazy_import()

        # 0. 准备输出路径
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.jpeg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        # 1. 确定图表类型
        resolved_type = self._resolve_chart_type(description, chart_type)

        # 2. 创建图形并绘制
        try:
            fig, ax = self._plt.subplots(figsize=self._figsize)
        except Exception as e:
            raise RuntimeError(f"创建 Matplotlib 图形失败: {e}") from e

        try:
            # 3. 按类型分发绘制
            plotter = self._get_plotter(resolved_type)
            plotter(ax, description, data or {})

            # 4. 应用 Nature 样式
            self.nature_style(ax)

            # 5. 处理中文标签
            self._tag_chinese_support(ax)

            # 6. 保存
            fig.savefig(str(out), dpi=self._dpi, facecolor='white',
                        edgecolor='none')
            logger.info("图表已生成: %s (类型=%s, 描述='%s')",
                        out, resolved_type, description[:60])
        except Exception as e:
            raise RuntimeError(f"绘制图表失败 ({resolved_type}): {e}") from e
        finally:
            self._plt.close(fig)

        return str(out)

    # ── 图表类型解析 ─────────────────────────────────────────

    def _resolve_chart_type(self, description: str, preferred: str) -> str:
        """解析图表类型。

        如果 preferred 为 'auto'，则根据描述关键词自动推断。
        """
        if preferred and preferred != 'auto':
            if preferred in self.CHART_TYPES:
                return preferred
            raise ValueError(
                f"不支持的图表类型 '{preferred}'。"
                f"可用类型: {', '.join(self.CHART_TYPES.keys())}"
            )

        desc = description.lower()

        # 关键词 → 类型映射
        patterns = [
            (r'频谱|频率.*幅值|频域|fft|功率谱', 'line'),
            (r'对比|对比图|多线|多条曲线|multiline', 'multiline'),
            (r'柱状|bar|直方', 'bar'),
            (r'分组.*柱|grouped.*bar', 'grouped_bar'),
            (r'散点|scatter|分布.*点', 'scatter'),
            (r'直方图|分布|histogram|概率', 'histogram'),
            (r'箱线|boxplot|盒须', 'boxplot'),
            (r'热力|heatmap|相关.*矩阵', 'heatmap'),
        ]

        for pattern, ctype in patterns:
            if re.search(pattern, desc):
                return ctype

        # 检测对比关键词暗示 multiline
        if re.search(r'原始.*?滤波|对比|before.*after', desc):
            return 'multiline'

        return 'line'  # 默认折线图

    def _get_plotter(self, chart_type: str):
        """根据图表类型获取对应的绘制函数。"""
        plotters = {
            'line': self._plot_line,
            'bar': self._plot_bar,
            'scatter': self._plot_scatter,
            'histogram': self._plot_histogram,
            'boxplot': self._plot_boxplot,
            'heatmap': self._plot_heatmap,
            'multiline': self._plot_multiline,
            'grouped_bar': self._plot_grouped_bar,
        }
        plotter = plotters.get(chart_type)
        if plotter is None:
            raise ValueError(f"图表类型 '{chart_type}' 没有对应的绘制函数")
        return plotter

    # ── 中文字体处理 ─────────────────────────────────────────

    def _tag_chinese_support(self, ax) -> None:
        """为所有文本元素设置 Fallback 字体，确保中文正常显示。

        修改 ax 上的 title、xlabel、ylabel、tick labels、legend 等
        文本对象，让它们用系统中文字体渲染。
        """
        cjk = _ensure_cjk_font()
        for item in (
                [ax.title, ax.xaxis.label, ax.yaxis.label]
                + ax.get_xticklabels() + ax.get_yticklabels()
                + (ax.get_legend().get_texts() if ax.get_legend() else [])
        ):
            try:
                item.set_fontproperties(self._plt.matplotlib.font_manager
                                        .FontProperties(family=[cjk, 'Arial']))
            except Exception:
                pass

    # ── 数据生成（从描述中提取或构造示例数据）──────────────

    @staticmethod
    def _extract_numbers(text: str, count: int = 10, default_max: float = 1.0) -> list[float]:
        """从文本中提取数值列表，若不足则用随机数补充。"""
        import random
        import math

        numbers = [float(n) for n in re.findall(r'\d+\.?\d*', text)]

        if len(numbers) >= count:
            return numbers[:count]

        # 补充随机数
        seed = hash(text) & 0xFFFFFFFF
        rng = random.Random(seed)
        needed = count - len(numbers)
        numbers.extend(
            [rng.random() * default_max for _ in range(needed)]
        )
        return numbers

    @staticmethod
    def _extract_axis_labels(description: str) -> tuple[str, str]:
        """从描述中提取 x/y 轴标签。"""
        xlabel = ''
        ylabel = ''

        # 匹配常见模式: "x轴/x坐标为...","y轴/y坐标为..."
        m = re.search(r'[xX][轴|坐标|轴为|标签]\s*[：:]\s*(.+?)(?:[，,。;]|$)', description)
        if m:
            xlabel = m.group(1).strip()

        m = re.search(r'[yY][轴|坐标|轴为|标签]\s*[：:]\s*(.+?)(?:[，,。;]|$)', description)
        if m:
            ylabel = m.group(1).strip()

        # 更自由的模式："*为x轴、*为y轴"
        m = re.search(r'(.+?)(?:\s*为\s*[xX]轴|x轴标签)[，,。\s]*(.+?)(?:\s*为\s*[yY]轴|y轴标签)', description)
        if m and not xlabel:
            xlabel = m.group(1).strip()
        if m and not ylabel:
            ylabel = m.group(2).strip()

        if not xlabel:
            m = re.search(r'[xX][轴|坐标][^，,。]*?[是为](.+?)(?:[，,。]|$)', description)
            if m:
                xlabel = m.group(1).strip()
        if not ylabel:
            m = re.search(r'[yY][轴|坐标][^，,。]*?[是为](.+?)(?:[，,。]|$)', description)
            if m:
                ylabel = m.group(1).strip()

        return xlabel or 'x', ylabel or 'y'

    def _extract_title(self, description: str) -> str:
        """从描述中提取图表标题（第一句或 '绘制*' 后的内容）。"""
        # 去掉"绘制/生成/画"等前缀
        cleaned = re.sub(r'^(请|帮我)?(绘制|生成|画|画出|画一个|制作)', '', description).strip()
        # 取第一句
        title = cleaned.split('。')[0].split('.')[0].strip()
        # 去掉末尾标点
        title = title.rstrip('，,。.')
        return title[:80] if title else '科学图表'

    # ── 各类型绘制函数 ───────────────────────────────────────

    def _plot_line(self, ax, description: str, data: dict):
        """绘制折线图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'x' in data and 'y' in data:
            x_vals = data['x']
            y_vals = data['y']
            label = data.get('label', 'data')
        else:
            numbers = self._extract_numbers(description, 20, 5.0)
            x_vals = list(range(len(numbers)))
            y_vals = numbers

        colors = self.nature_colors(1)
        ax.plot(x_vals, y_vals, color=colors[0], linewidth=0.8,
                marker='o', markersize=2.5, markerfacecolor=colors[0],
                markeredgewidth=0)

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_bar(self, ax, description: str, data: dict):
        """绘制柱状图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'categories' in data and 'values' in data:
            categories = data['categories']
            values = data['values']
        else:
            categories = [f'组{i+1}' for i in range(6)]
            values = self._extract_numbers(description, 6, 3.0)

        colors = self.nature_colors(1)
        ax.bar(categories, values, color=colors[0], width=0.6,
               edgecolor='white', linewidth=0.3)

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_scatter(self, ax, description: str, data: dict):
        """绘制散点图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'x' in data and 'y' in data:
            x_vals = data['x']
            y_vals = data['y']
        else:
            import random
            seed = hash(description) & 0xFFFFFFFF
            rng = random.Random(seed)
            n = 30
            x_vals = [rng.random() * 10 for _ in range(n)]
            y_vals = [v + rng.gauss(0, 0.5) for v in x_vals]

        colors = self.nature_colors(1)
        ax.scatter(x_vals, y_vals, color=colors[0], s=8,
                   edgecolors='none', alpha=0.7)

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_histogram(self, ax, description: str, data: dict):
        """绘制直方图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'values' in data:
            values = data['values']
        else:
            import random
            seed = hash(description) & 0xFFFFFFFF
            rng = random.Random(seed)
            values = [rng.gauss(0, 1) for _ in range(200)]

        colors = self.nature_colors(1)
        ax.hist(values, bins='auto', color=colors[0], alpha=0.7,
                edgecolor='white', linewidth=0.3)

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel or '频数')
        ax.set_title(title)

    def _plot_boxplot(self, ax, description: str, data: dict):
        """绘制箱线图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'groups' in data:
            groups_data = data['groups']
            labels = data.get('labels', [f'组{i+1}' for i in range(len(groups_data))])
        else:
            import random
            seed = hash(description) & 0xFFFFFFFF
            rng = random.Random(seed)
            labels = [f'组{i+1}' for i in range(4)]
            groups_data = []
            for i in range(4):
                mean = i * 0.5
                groups_data.append([rng.gauss(mean, 0.3) for _ in range(30)])

        colors = self.nature_colors(1)
        bp = ax.boxplot(groups_data, labels=labels, patch_artist=True,
                        widths=0.5, showmeans=True,
                        meanprops=dict(marker='D', markerfacecolor='white',
                                       markeredgecolor='black', markersize=3))
        for patch in bp['boxes']:
            patch.set_facecolor(colors[0])
            patch.set_alpha(0.7)

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_heatmap(self, ax, description: str, data: dict):
        """绘制热力图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        if 'matrix' in data:
            matrix = data['matrix']
            row_labels = data.get('row_labels', None)
            col_labels = data.get('col_labels', None)
        else:
            import random
            seed = hash(description) & 0xFFFFFFFF
            rng = random.Random(seed)
            n = 6
            matrix = [[rng.random() for _ in range(n)] for _ in range(n)]
            row_labels = [f'R{i+1}' for i in range(n)]
            col_labels = [f'C{i+1}' for i in range(n)]

        im = ax.imshow(matrix, cmap='Blues', aspect='auto',
                       interpolation='nearest')

        # 颜色条
        cbar = self._plt.colorbar(im, ax=ax, shrink=0.75, pad=0.02)
        cbar.ax.tick_params(labelsize=6)

        # 坐标轴标签
        if row_labels:
            ax.set_yticks(range(len(row_labels)))
            ax.set_yticklabels(row_labels, fontsize=6)
        if col_labels:
            ax.set_xticks(range(len(col_labels)))
            ax.set_xticklabels(col_labels, fontsize=6, rotation=30, ha='right')

        # 显示数值
        for i in range(len(matrix)):
            for j in range(len(matrix[0])):
                val = matrix[i][j]
                ax.text(j, i, f'{val:.2f}', ha='center', va='center',
                        fontsize=5, color='white' if val > 0.5 else 'black')

        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_multiline(self, ax, description: str, data: dict):
        """绘制多线对比图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        n_lines = 3
        colors = self.nature_colors(n_lines)

        # 从描述中提取行名
        line_names = ['原始信号', '滤波后信号', '参考信号']
        description_lower = description.lower()
        if '原始' in description_lower and '滤波' in description_lower:
            line_names = ['原始信号', '滤波后信号']
        elif '对比' in description_lower:
            line_names = ['方法A', '方法B', '方法C']

        if 'x' in data and isinstance(data.get('y'), list):
            x_vals = data['x']
            y_sets = data['y']
        else:
            import math
            n_pts = 200
            x_vals = [i / n_pts * 2 * math.pi for i in range(n_pts)]
            seed = hash(description) & 0xFFFFFFFF
            rng = __import__('random').Random(seed)
            y_sets = []
            for li in range(len(line_names)):
                base = [math.sin(x + li * 0.5) for x in x_vals]
                noise = [rng.gauss(0, 0.05) for _ in range(n_pts)]
                y_sets.append([b + n for b, n in zip(base, noise)])

        for i, y_vals in enumerate(y_sets):
            color = colors[i % len(colors)]
            label = data.get('labels', line_names)[i] if data.get('labels') else line_names[i]
            ax.plot(x_vals, y_vals, color=color, linewidth=0.7,
                    label=label)

        ax.legend(frameon=False, loc='best')
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    def _plot_grouped_bar(self, ax, description: str, data: dict):
        """绘制分组柱状图。"""
        xlabel, ylabel = self._extract_axis_labels(description)
        title = self._extract_title(description)

        groups = 4
        n_bars = 3
        colors = self.nature_colors(n_bars)

        if 'categories' in data and 'groups' in data:
            categories = data['categories']
            groups_data = data['groups']
            groups = len(categories)
            n_bars = len(groups_data)
        else:
            import random
            seed = hash(description) & 0xFFFFFFFF
            rng = random.Random(seed)
            categories = [f'类别{i+1}' for i in range(groups)]
            group_labels = data.get('group_labels', [f'组{j+1}' for j in range(n_bars)])
            groups_data = [[rng.random() * 2 for _ in range(groups)] for _ in range(n_bars)]

        group_labels = data.get('group_labels',
                                [f'组{j+1}' for j in range(len(groups_data))])

        x_pos = list(range(groups))
        width = 0.6 / n_bars

        for i in range(n_bars):
            offset = (i - (n_bars - 1) / 2) * width
            ax.bar([p + offset for p in x_pos], groups_data[i],
                   width=width, color=colors[i % len(colors)],
                   label=group_labels[i], edgecolor='white', linewidth=0.3)

        ax.set_xticks(x_pos)
        ax.set_xticklabels(categories)
        ax.legend(frameon=False, loc='best')
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)

    # ── 预置图表模板 ─────────────────────────────────────────

    def render_spectrum(self, output_path: str,
                        frequencies: Optional[list[float]] = None,
                        magnitudes: Optional[list[float]] = None,
                        title: str = '频谱图',
                        fs: float = 10000) -> str:
        """频谱图模板 — frequency vs magnitude，适合信号处理实验报告。

        Args:
            output_path: 输出路径
            frequencies: 频率数组（Hz），None 则自动生成
            magnitudes: 幅值数组，None 则自动生成
            title: 图表标题
            fs: 采样频率（Hz），用于自动生成数据

        Returns:
            图片路径
        """
        self._lazy_import()

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        fig, ax = self._plt.subplots(figsize=self._figsize)

        try:
            import math
            import random
            rng = random.Random(42)

            if frequencies is None or magnitudes is None:
                n_pts = 512
                freqs = [fs / 2 * i / n_pts for i in range(n_pts // 2 + 1)]
                # 模拟两个频率峰值：500Hz 和 3000Hz
                mags = []
                for f in freqs:
                    base = 0.01 + rng.random() * 0.005
                    if 480 <= f <= 520:
                        base += 0.8 * math.exp(-((f - 500) ** 2) / (2 * 30 ** 2))
                    if 2950 <= f <= 3050:
                        base += 0.4 * math.exp(-((f - 3000) ** 2) / (2 * 40 ** 2))
                    mags.append(base)
            else:
                freqs = frequencies
                mags = magnitudes

            colors = self.nature_colors(1)
            ax.plot(freqs, mags, color=colors[0], linewidth=0.7)
            ax.set_xlabel('频率 (Hz)')
            ax.set_ylabel('幅值')
            ax.set_title(title)

            self.nature_style(ax)
            fig.savefig(str(out), dpi=self._dpi, facecolor='white')
            logger.info("频谱图已生成: %s", out)
        except Exception as e:
            raise RuntimeError(f"生成频谱图失败: {e}") from e
        finally:
            self._plt.close(fig)

        return str(out)

    def render_filter_comparison(self, output_path: str,
                                  original: Optional[list[float]] = None,
                                  filtered: Optional[list[float]] = None,
                                  title: str = '滤波前后对比',
                                  fs: float = 1000) -> str:
        """滤波对比图模板 — 原始信号 vs 滤波后信号。

        Args:
            output_path: 输出路径
            original: 原始信号数据
            filtered: 滤波后信号数据
            title: 图表标题
            fs: 采样频率（Hz）

        Returns:
            图片路径
        """
        self._lazy_import()

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        fig, (ax1, ax2) = self._plt.subplots(2, 1, figsize=(3.5, 3.0),
                                              sharex=True)

        try:
            import math
            import random
            rng = random.Random(42)
            n = 200
            t = [i / fs for i in range(n)]

            if original is None:
                original = [math.sin(2 * math.pi * 5 * x)
                            + 0.3 * math.sin(2 * math.pi * 50 * x)
                            for x in t]

            if filtered is None:
                # 模拟低通滤波效果（去除高频成分）
                filtered = [math.sin(2 * math.pi * 5 * x)
                            + 0.05 * math.sin(2 * math.pi * 50 * x)
                            for x in t]

            colors = self.nature_colors(2)

            ax1.plot(t, original, color=colors[0], linewidth=0.7)
            ax1.set_ylabel('原始信号')
            ax1.set_title(title)

            ax2.plot(t, filtered, color=colors[1], linewidth=0.7)
            ax2.set_ylabel('滤波后信号')
            ax2.set_xlabel('时间 (s)')

            for ax in (ax1, ax2):
                self.nature_style(ax)

            fig.tight_layout(pad=0.8)
            fig.savefig(str(out), dpi=self._dpi, facecolor='white')
            logger.info("滤波对比图已生成: %s", out)
        except Exception as e:
            raise RuntimeError(f"生成滤波对比图失败: {e}") from e
        finally:
            self._plt.close(fig)

        return str(out)

    def render_experiment_data(self, output_path: str,
                                x_data: Optional[list[float]] = None,
                                y_data: Optional[list[float]] = None,
                                y_err: Optional[list[float]] = None,
                                xlabel: str = '参数',
                                ylabel: str = '测量值',
                                title: str = '实验数据') -> str:
        """实验数据图模板 — 带误差棒的折线图。

        Args:
            output_path: 输出路径
            x_data: x 轴数据
            y_data: y 轴数据
            y_err: 误差棒数据
            xlabel: x 轴标签
            ylabel: y 轴标签
            title: 图表标题

        Returns:
            图片路径
        """
        self._lazy_import()

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        fig, ax = self._plt.subplots(figsize=self._figsize)

        try:
            import random
            rng = random.Random(42)

            if x_data is None:
                x_data = [i for i in range(1, 9)]
            if y_data is None:
                y_data = [v + rng.random() * 0.5 for v in [1.0, 1.8, 2.5, 3.0,
                                                           3.2, 3.5, 3.8, 3.9]]
            if y_err is None:
                y_err = [0.15 + rng.random() * 0.1 for _ in y_data]

            colors = self.nature_colors(1)
            ax.errorbar(x_data, y_data, yerr=y_err, color=colors[0],
                        fmt='o-', capsize=2, capthick=0.5,
                        linewidth=0.8, markersize=3,
                        markeredgewidth=0)

            ax.set_xlabel(xlabel)
            ax.set_ylabel(ylabel)
            ax.set_title(title)

            self.nature_style(ax)
            fig.savefig(str(out), dpi=self._dpi, facecolor='white')
            logger.info("实验数据图已生成: %s", out)
        except Exception as e:
            raise RuntimeError(f"生成实验数据图失败: {e}") from e
        finally:
            self._plt.close(fig)

        return str(out)

    def render_model_comparison(self, output_path: str,
                                 model_names: Optional[list[str]] = None,
                                 metrics: Optional[list[float]] = None,
                                 title: str = '模型性能对比',
                                 ylabel: str = '准确率 (%)') -> str:
        """模型性能对比模板 — 多模型指标对比柱状图。

        Args:
            output_path: 输出路径
            model_names: 模型名称列表
            metrics: 各模型指标值
            title: 图表标题
            ylabel: y 轴标签

        Returns:
            图片路径
        """
        self._lazy_import()

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        fig, ax = self._plt.subplots(figsize=self._figsize)

        try:
            import random
            rng = random.Random(42)

            if model_names is None:
                model_names = ['CNN', 'RNN', 'ResNet', 'Transformer', 'ViT']
            if metrics is None:
                metrics = [rng.uniform(70, 98) for _ in model_names]

            colors = self.nature_colors(len(model_names))

            bars = ax.bar(model_names, metrics, color=colors[:len(model_names)],
                          width=0.55, edgecolor='white', linewidth=0.3)

            # 在柱子上方标注数值
            for bar, val in zip(bars, metrics):
                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                        f'{val:.1f}', ha='center', va='bottom', fontsize=6)

            ax.set_ylabel(ylabel)
            ax.set_title(title)

            self.nature_style(ax)
            fig.savefig(str(out), dpi=self._dpi, facecolor='white')
            logger.info("模型对比图已生成: %s", out)
        except Exception as e:
            raise RuntimeError(f"生成模型对比图失败: {e}") from e
        finally:
            self._plt.close(fig)

        return str(out)

    def render_flow_diagram(self, output_path: str,
                             steps: Optional[list[str]] = None,
                             title: str = '流程图') -> str:
        """流程示意图模板 — 用 matplotlib patches 绘制的流程框图。

        Args:
            output_path: 输出路径
            steps: 流程步骤名称列表
            title: 图表标题

        Returns:
            图片路径
        """
        self._lazy_import()

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.suffix.lower() not in ('.png', '.jpg', '.tiff', '.pdf'):
            out = out.with_suffix('.png')

        from matplotlib.patches import FancyBboxPatch

        if steps is None:
            steps = ['数据采集', '预处理', '特征提取',
                     '模型训练', '评估验证', '结果分析']

        fig, ax = self._plt.subplots(figsize=(self._figsize[0] * 1.2,
                                               self._figsize[1] * 0.8))
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
        ax.set_title(title, fontsize=9, pad=10)

        n = len(steps)
        colors = self.nature_colors(n)

        # 自动布局：根据步数自适应
        if n <= 4:
            # 水平排列
            box_w = 0.8 / n
            box_h = 0.25
            y_center = 0.5
            x_positions = [0.1 + i * (0.8 / (n - 1 or 1)) for i in range(n)]

            for i, (step, color) in enumerate(zip(steps, colors)):
                x = x_positions[i] - box_w / 2
                rect = FancyBboxPatch(
                    (x, y_center - box_h / 2), box_w, box_h,
                    boxstyle=f"round,pad=0.03",
                    facecolor=color, edgecolor='white', linewidth=0.5,
                    alpha=0.85,
                )
                ax.add_patch(rect)
                ax.text(x + box_w / 2, y_center, step,
                        ha='center', va='center', fontsize=6.5,
                        color='white', fontweight='bold')

                # 箭头
                if i < n - 1:
                    ax.annotate('', xy=(x_positions[i + 1] - box_w / 2 - 0.01, y_center),
                                xytext=(x_positions[i] + box_w / 2 + 0.01, y_center),
                                arrowprops=dict(arrowstyle='->', color='#888',
                                                lw=0.8))
        else:
            # 蛇形排列（多行）
            n_cols = min(4, n)
            n_rows = (n + n_cols - 1) // n_cols
            box_w = 0.7 / n_cols
            box_h = 0.6 / n_rows

            for i, (step, color) in enumerate(zip(steps, colors)):
                row = i // n_cols
                col = i % n_cols
                x = 0.1 + col * (box_w + 0.05)
                y = 0.85 - row * (box_h + 0.1)

                rect = FancyBboxPatch(
                    (x, y - box_h / 2), box_w, box_h,
                    boxstyle=f"round,pad=0.03",
                    facecolor=color, edgecolor='white', linewidth=0.5,
                    alpha=0.85,
                )
                ax.add_patch(rect)
                ax.text(x + box_w / 2, y, step,
                        ha='center', va='center', fontsize=6,
                        color='white', fontweight='bold',
                        wrap=True)

                # 箭头
                if i < n - 1:
                    next_row = (i + 1) // n_cols
                    next_col = (i + 1) % n_cols
                    nx = 0.1 + next_col * (box_w + 0.05)
                    ny = 0.85 - next_row * (box_h + 0.1)

                    if next_row == row:
                        # 水平向右
                        start_x = x + box_w + 0.005
                        end_x = nx - 0.005
                        ax.annotate('', xy=(end_x, y),
                                    xytext=(start_x, y),
                                    arrowprops=dict(arrowstyle='->', color='#888',
                                                    lw=0.6))
                    else:
                        # 换行：先向下再向左
                        mid_y = y - box_h / 2 - 0.02
                        ax.plot([x + box_w / 2, x + box_w / 2],
                                [y - box_h / 2, mid_y],
                                color='#888', lw=0.6)
                        ax.plot([x + box_w / 2, nx + box_w / 2],
                                [mid_y, mid_y],
                                color='#888', lw=0.6)
                        ax.annotate('', xy=(nx + box_w / 2, ny + box_h / 2 + 0.005),
                                    xytext=(nx + box_w / 2, mid_y),
                                    arrowprops=dict(arrowstyle='->', color='#888',
                                                    lw=0.6))

        fig.savefig(str(out), dpi=self._dpi, facecolor='white',
                    bbox_inches='tight')
        logger.info("流程图已生成: %s", out)
        self._plt.close(fig)

        return str(out)

    # ── CSV 数据解析 ─────────────────────────────────────────

    @staticmethod
    def parse_csv_data(csv_text: str) -> dict:
        """将 CSV 格式文本解析为数据字典。

        支持格式：
        - 逗号/制表符分隔
        - 首行为列名
        - 数值列自动转为 float

        Args:
            csv_text: CSV 格式的文本数据

        Returns:
            {column_name: [values], ...}
            or {'x': [...], 'y': [...]} 如果只有两列
        """
        if not csv_text or not csv_text.strip():
            return {}

        lines = [l.strip() for l in csv_text.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            return {}

        # 检测分隔符（优先制表符，其次逗号）
        tab_count = lines[0].count('\t')
        comma_count = lines[0].count(',')
        delimiter = '\t' if tab_count >= comma_count else ','

        header = [h.strip().strip('"\'') for h in lines[0].split(delimiter)]
        data: dict = {h: [] for h in header}

        for line in lines[1:]:
            values = line.split(delimiter)
            for i, val in enumerate(values):
                if i >= len(header):
                    break
                val = val.strip().strip('"\'')
                try:
                    data[header[i]].append(float(val))
                except ValueError:
                    data[header[i]].append(val)

        # 简化：如果只有两列，命名为 x, y
        if len(header) == 2:
            return {'x': data[header[0]], 'y': data[header[1]],
                    '_labels': header}

        return data

    # ── 渲染为 bytes（用于预览）──────────────────────────────

    def render_to_bytes(self, description: str,
                        chart_type: str = 'auto',
                        data: Optional[dict] = None) -> bytes:
        """渲染图表为 PNG bytes（用于 GUI 预览，不写磁盘）。

        Args:
            description: 图表描述
            chart_type: 图表类型
            data: 结构化数据

        Returns:
            PNG 图片的 bytes
        """
        self._lazy_import()

        resolved_type = self._resolve_chart_type(description, chart_type)

        fig, ax = self._plt.subplots(figsize=self._figsize)
        try:
            plotter = self._get_plotter(resolved_type)
            plotter(ax, description, data or {})
            self.nature_style(ax)
            self._tag_chinese_support(ax)

            buf = io.BytesIO()
            fig.savefig(buf, format='png', dpi=150, facecolor='white',
                        edgecolor='none')
            buf.seek(0)
            return buf.getvalue()
        finally:
            self._plt.close(fig)

    # ── 插入到 DOCX ──────────────────────────────────────────

    def insert_into_docx(self, docx_path: str, image_bytes: bytes,
                         output_path: Optional[str] = None) -> str:
        """将图表图片插入到 DOCX 文档末尾。

        Args:
            docx_path: 目标 .docx 路径
            image_bytes: PNG 图片 bytes
            output_path: 输出路径，None 则自动生成

        Returns:
            保存后的 .docx 路径
        """
        from docx import Document as DocxDoc
        from docx.shared import Inches
        import tempfile

        doc = DocxDoc(docx_path)
        # 将 bytes 写入临时文件
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name

        try:
            # 在文档末尾添加新段落并插入图片
            doc.add_paragraph()  # 空行
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(tmp_path, width=Inches(5.0))
            doc.add_paragraph()  # 空行

            out_path = output_path or str(
                Path(docx_path).parent / f"{Path(docx_path).stem}_图表版.docx"
            )
            doc.save(out_path)
            logger.info("图表已插入文档: %s", out_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

        return out_path
